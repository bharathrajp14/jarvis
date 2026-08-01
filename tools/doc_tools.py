# tools/doc_tools.py — BR JARVIS Executive Document Generator Engine v2
"""
Automated Executive Document Creator for Microsoft Word (.docx), PDF (.pdf), HTML (.html), and Markdown (.md).
Supports Cover Pages, Styled Tables, Callout Boxes, Code Syntax Blocks, Headers, Footers, and Auto-Launching.
Includes robust method alias protection on Document instances (addpagebreak, add_pagebreak) to prevent AI syntax errors.
"""
from __future__ import annotations

import json
import os
import re
import sys
import subprocess
import time
from pathlib import Path
from typing import Any

from tools.registry import register_tool

try:
    import docx  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False


def _get_workspace_dir() -> Path:
    return Path(__file__).resolve().parent.parent


# ── XML Styling Helpers for DOCX ──────────────────────────────────────────
def set_cell_background(cell: Any, fill_hex: str):
    """Set the background color of a table cell in hex format (e.g. 'F2F4F7')."""
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_left_border(cell: Any, border_hex: str = '1B365D', border_size_pt: float = 4.0):
    """Set a thick left border and remove top, bottom, and right borders for callout boxes."""
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(int(border_size_pt * 8)))
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_hex)
    tcBorders.append(left)
    
    for side in ('top', 'bottom', 'right'):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
        
    tcPr.append(tcBorders)


def set_cell_margins(cell: Any, top_pt: int = 8, bottom_pt: int = 8, left_pt: int = 12, right_pt: int = 12):
    """Set inner padding margins for table cells."""
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top_pt * 14), ('bottom', bottom_pt * 14), ('left', left_pt * 14), ('right', right_pt * 14)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_docx_callout(doc: Any, text: str, border_hex: str = '1B365D', fill_hex: str = 'F2F4F7'):
    """Add a styled callout box with a gray background and thick navy left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    
    set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=14, right_pt=14)
    set_cell_background(cell, fill_hex)
    set_cell_left_border(cell, border_hex, 4.0)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def inject_document_aliases(doc: Any):
    """Inject robust method aliases onto docx Document instance to prevent AI NameError crashes."""
    doc.addpagebreak = doc.add_page_break
    doc.add_pagebreak = doc.add_page_break
    doc.pagebreak = doc.add_page_break
    doc.addcallout = lambda text: add_docx_callout(doc, text)
    doc.add_callout = lambda text: add_docx_callout(doc, text)


def _add_paragraph_runs(p: Any, text: str, font_name: str = 'Calibri', font_size_pt: float = 11.0, font_color_rgb=(0x22, 0x22, 0x22)):
    """Split string on markdown '**' bold and '*' or '_' italic markers and append runs to paragraph."""
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            subparts = part.split("*")
            is_italic = False
            for subpart in subparts:
                if subpart:
                    run = p.add_run(subpart)
                    run.bold = is_bold
                    run.italic = is_italic
                    run.font.name = font_name
                    run.font.size = Pt(font_size_pt)
                    if font_color_rgb:
                        run.font.color.rgb = RGBColor(*font_color_rgb)
                is_italic = not is_italic
        is_bold = not is_bold


# ── Advanced DOCX Document Builder ─────────────────────────────────────────
def _build_executive_docx(title: str, subtitle: str, author: str, content: str, out_path: Path, cover_page: bool = True) -> str:
    """Build a publication-grade Word Document with Cover Page, Callouts, Tables, and Code Blocks."""
    doc = docx.Document()
    inject_document_aliases(doc)
    
    # Configure Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- COVER PAGE ---
    if cover_page:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(80)
        title_p.paragraph_format.space_after = Pt(12)
        
        t_run = title_p.add_run(title.upper())
        t_run.font.name = 'Calibri'
        t_run.font.size = Pt(32)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        if subtitle:
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_p.paragraph_format.space_after = Pt(100)
            s_run = sub_p.add_run(subtitle)
            s_run.font.name = 'Calibri'
            s_run.font.size = Pt(14)
            s_run.font.italic = True
            s_run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
            
        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_p.paragraph_format.space_after = Pt(4)
        a_run = auth_p.add_run(f"Authored by: {author or 'BR JARVIS Systems Engine'}")
        a_run.font.name = 'Calibri'
        a_run.font.size = Pt(11)
        a_run.font.bold = True
        a_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d_run = date_p.add_run(time.strftime("%B %Y Edition"))
        d_run.font.name = 'Calibri'
        d_run.font.size = Pt(10)
        d_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        
        doc.add_page_break()
    else:
        h1 = doc.add_heading(title, level=0)
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if subtitle:
            p_sub = doc.add_paragraph()
            r_sub = p_sub.add_run(subtitle)
            r_sub.font.italic = True
            r_sub.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    # --- BODY CONTENT PARSER ---
    lines = content.splitlines()
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        line_s = line.strip()

        # Handle Code Block Start / End
        if line_s.startswith("```"):
            if in_code_block:
                # Flush code block
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                c = tbl.cell(0, 0)
                c.width = Inches(6.0)
                set_cell_background(c, '1E293B')  # Slate Dark
                set_cell_margins(c, top_pt=6, bottom_pt=6, left_pt=10, right_pt=10)
                p_code = c.paragraphs[0]
                p_code.paragraph_format.space_before = Pt(2)
                p_code.paragraph_format.space_after = Pt(2)
                p_code.paragraph_format.line_spacing = 1.0
                r_code = p_code.add_run("\n".join(code_lines))
                r_code.font.name = 'Consolas'
                r_code.font.size = Pt(9.5)
                r_code.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)
                
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not line_s:
            i += 1
            continue

        # Handle Callout Box (> [!NOTE] or > Callout)
        if line_s.startswith(">"):
            callout_text = line_s.lstrip("> ").strip()
            if callout_text.startswith("[!"):
                callout_text = callout_text.split("]", 1)[-1].strip()
            add_docx_callout(doc, callout_text)
            i += 1
            continue

        # Handle Markdown Tables
        if line_s.startswith("|") and i + 1 < len(lines) and lines[i+1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            parsed_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                is_sep = all(all(ch in "-: " for ch in cell) for cell in cells) if cells else False
                if not is_sep:
                    parsed_rows.append(cells)

            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                tbl = doc.add_table(rows=0, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r_idx, row_cells in enumerate(parsed_rows):
                    row = tbl.add_row()
                    for c_idx, cell_text in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
                            p_cell = cell.paragraphs[0]
                            p_cell.paragraph_format.space_before = Pt(2)
                            p_cell.paragraph_format.space_after = Pt(2)
                            
                            if r_idx == 0:
                                set_cell_background(cell, '1B365D')  # Navy Header
                                _add_paragraph_runs(p_cell, cell_text, font_color_rgb=(0xFF, 0xFF, 0xFF))
                                for run in p_cell.runs:
                                    run.bold = True
                            else:
                                bg_fill = 'F9FAFB' if r_idx % 2 == 1 else 'FFFFFF'
                                set_cell_background(cell, bg_fill)
                                _add_paragraph_runs(p_cell, cell_text, font_color_rgb=(0x22, 0x22, 0x22))
                                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # Headings
        if line_s.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[2:])
            r.font.name = 'Calibri'
            r.font.size = Pt(18)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        elif line_s.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[3:])
            r.font.name = 'Calibri'
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x00, 0x80, 0x80)  # Teal Accent
        elif line_s.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[4:])
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        # Bullet list items
        elif line_s.startswith("- ") or line_s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            _add_paragraph_runs(p, line_s[2:])
        # Numbered list items
        elif re.match(r'^\d+\.\s+', line_s):
            match = re.match(r'^(\d+\.\s+)(.*)', line_s)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run_num = p.add_run(match.group(1))
            run_num.font.name = 'Calibri'
            run_num.font.bold = True
            run_num.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            _add_paragraph_runs(p, match.group(2))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            _add_paragraph_runs(p, line_s)

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return str(out_path)


def _pdf_cell(pdf, w, h, text="", ln=False, align="L"):
    """Helper for fpdf2 / fpdf compatibility."""
    try:
        if ln:
            pdf.cell(w, h, text=text, align=align, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.cell(w, h, text=text, align=align)
    except Exception:
        try:
            pdf.cell(w, h, txt=text, ln=ln, align=align)
        except Exception:
            pdf.cell(w, h, text)

# ── Advanced PDF Builder ───────────────────────────────────────────────────
def _build_executive_pdf(title: str, subtitle: str, author: str, content: str, out_path: Path) -> str:
    """Build a styled PDF document using FPDF."""
    pdf = FPDF()
    pdf.add_page()
    
    # Title
    pdf.set_font("Helvetica", size=18, style="B")
    _pdf_cell(pdf, 0, 10, text=title.encode("latin-1", "replace").decode("latin-1"), ln=True, align="L")
    
    if subtitle:
        pdf.set_font("Helvetica", size=11, style="I")
        pdf.set_text_color(100, 110, 120)
        _pdf_cell(pdf, 0, 8, text=subtitle.encode("latin-1", "replace").decode("latin-1"), ln=True, align="L")
        pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    lines = content.splitlines()
    for line in lines:
        line_clean = line.encode("latin-1", "replace").decode("latin-1").strip()
        if not line_clean:
            pdf.ln(3)
            continue

        if line_clean.startswith("# "):
            pdf.ln(4)
            pdf.set_font("Helvetica", size=14, style="B")
            _pdf_cell(pdf, 0, 8, text=line_clean[2:], ln=True)
            pdf.ln(2)
        elif line_clean.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", size=12, style="B")
            _pdf_cell(pdf, 0, 7, text=line_clean[3:], ln=True)
            pdf.ln(2)
        elif line_clean.startswith("- ") or line_clean.startswith("* "):
            pdf.set_font("Helvetica", size=10)
            pdf.write(5, "  *  " + line_clean[2:])
            pdf.ln(6)
        else:
            pdf.set_font("Helvetica", size=10)
            pdf.write(5, line_clean)
            pdf.ln(6)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))
    return str(out_path)


# ── Advanced Glassmorphism HTML Builder ────────────────────────────────────
def _build_executive_html(title: str, subtitle: str, author: str, content: str, out_path: Path) -> str:
    """Build a modern responsive HTML document."""
    html_lines = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "<meta charset='UTF-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        f"<title>{title}</title>",
        "<style>",
        "  body { font-family: 'Inter', system-ui, -apple-system, sans-serif; background: #0f172a; color: #f8fafc; margin: 0; padding: 40px 20px; line-height: 1.6; }",
        "  .container { max-width: 900px; margin: 0 auto; background: rgba(30, 41, 59, 0.7); backdrop-filter: blur(16px); border: 1px solid rgba(255,255,255,0.1); border-radius: 16px; padding: 48px; box-shadow: 0 20px 40px rgba(0,0,0,0.5); }",
        "  h1 { font-size: 2.5rem; color: #38bdf8; margin-top: 0; margin-bottom: 8px; font-weight: 800; letter-spacing: -0.025em; }",
        "  .subtitle { font-size: 1.15rem; color: #94a3b8; font-style: italic; margin-bottom: 32px; border-bottom: 1px solid rgba(255,255,255,0.1); padding-bottom: 16px; }",
        "  h2 { color: #818cf8; font-size: 1.5rem; margin-top: 32px; border-bottom: 1px solid rgba(255,255,255,0.05); padding-bottom: 8px; }",
        "  h3 { color: #cbd5e1; font-size: 1.2rem; margin-top: 24px; }",
        "  p, li { color: #e2e8f0; font-size: 1.05rem; }",
        "  ul, ol { padding-left: 24px; }",
        "  .callout { background: rgba(56, 189, 248, 0.1); border-left: 4px solid #38bdf8; padding: 16px 20px; border-radius: 4px 8px 8px 4px; margin: 20px 0; font-style: italic; color: #e0f2fe; }",
        "  pre { background: #090d16; border: 1px solid #334155; padding: 16px; border-radius: 8px; overflow-x: auto; font-family: 'Consolas', monospace; color: #f8fafc; }",
        "  table { width: 100%; border-collapse: collapse; margin: 24px 0; font-size: 0.95rem; }",
        "  th { background: #1e1b4b; color: #a5b4fc; padding: 12px 16px; text-align: left; border-bottom: 2px solid #4338ca; }",
        "  td { padding: 12px 16px; border-bottom: 1px solid #334155; }",
        "  tr:nth-child(even) { background: rgba(255,255,255,0.02); }",
        "</style>",
        "</head>",
        "<body>",
        "<div class='container'>",
        f"<h1>{title}</h1>",
        f"<div class='subtitle'>{subtitle or 'Official Executive Document'} | {author or 'BR JARVIS'}</div>",
    ]

    for line in content.splitlines():
        line_s = line.strip()
        if not line_s:
            continue
        if line_s.startswith("# "):
            html_lines.append(f"<h2>{line_s[2:]}</h2>")
        elif line_s.startswith("## "):
            html_lines.append(f"<h2>{line_s[3:]}</h2>")
        elif line_s.startswith("### "):
            html_lines.append(f"<h3>{line_s[4:]}</h3>")
        elif line_s.startswith(">"):
            html_lines.append(f"<div class='callout'>{line_s.lstrip('> ').strip()}</div>")
        elif line_s.startswith("- ") or line_s.startswith("* "):
            html_lines.append(f"<ul><li>{line_s[2:]}</li></ul>")
        else:
            html_lines.append(f"<p>{line_s}</p>")

    html_lines.extend(["</div>", "</body>", "</html>"])
    
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(html_lines), encoding="utf-8")
    return str(out_path)


# ── Universal Document Creator Tool ────────────────────────────────────────
@register_tool(
    name="document_creator",
    description="Universal Executive Document Engine. Creates styled Word (.docx), PDF (.pdf), HTML (.html), and Markdown (.md) documents with Cover Pages, Styled Tables, Callouts, Code Blocks, and Auto-Launch.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title of the document"},
            "subtitle": {"type": "string", "description": "Subtitle or tagline for the document"},
            "author": {"type": "string", "description": "Author name (default: BR JARVIS AI)"},
            "content": {"type": "string", "description": "Main text content in structured Markdown (headings, bullets, tables, callouts)"},
            "filename": {"type": "string", "description": "Target filename or relative path (e.g., workspace/Books/Startup_Book.docx)"},
            "format": {"type": "string", "description": "Output format: docx | pdf | html | md (default: docx)"},
            "cover_page": {"type": "boolean", "description": "Whether to include an executive cover page (default: true)"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch the generated file (default: true)"}
        },
        "required": ["title", "content"]
    }
)
def document_creator(args: dict) -> str:
    """Universal Executive Document Engine."""
    title = args.get("title", "Document").strip()
    subtitle = args.get("subtitle", "").strip()
    author = args.get("author", "BR JARVIS Systems Engine").strip()
    content = args.get("content", "").strip()
    fmt = args.get("format", "docx").lower().strip()
    cover_page = args.get("cover_page", True)
    auto_open = args.get("auto_open", True)
    
    filename = args.get("filename", "").strip()
    if not filename:
        clean_title = re.sub(r'[^\w\-]', '_', title)
        filename = f"workspace/Documents/{clean_title}.{fmt}"
    elif not filename.endswith(f".{fmt}"):
        filename += f".{fmt}"
        
    out_path = _get_workspace_dir() / filename
    
    try:
        if fmt == "docx":
            if not _DOCX_AVAILABLE:
                return "Error: 'python-docx' library is not installed."
            saved_path = _build_executive_docx(title, subtitle, author, content, out_path, cover_page=cover_page)
        elif fmt == "pdf":
            if not _FPDF_AVAILABLE:
                return "Error: 'fpdf' library is not installed."
            saved_path = _build_executive_pdf(title, subtitle, author, content, out_path)
        elif fmt == "html":
            saved_path = _build_executive_html(title, subtitle, author, content, out_path)
        elif fmt == "md":
            out_path.parent.mkdir(parents=True, exist_ok=True)
            md_text = f"# {title}\n\n*{subtitle}*\n\n**Author:** {author}\n\n---\n\n{content}"
            out_path.write_text(md_text, encoding="utf-8")
            saved_path = str(out_path)
        else:
            return f"Error: Unsupported format '{fmt}'. Choose from docx, pdf, html, md."
    except Exception as e:
        return f"Error building document ({fmt}): {e}"
        
    if auto_open and sys.platform == "win32":
        try:
            subprocess.Popen(["cmd", "/c", "start", "", str(saved_path)], shell=False)
        except Exception:
            pass
            
    return f"⚡ Created Executive Document ({fmt.upper()}): '{saved_path}' and launched viewer."


# ── Backward Compatibility Tool Wrappers ──────────────────────────────────
@register_tool(
    name="create_word_document",
    description="Create a formatted Microsoft Word (.docx) document with cover page, headers, tables, callouts, and auto-launch.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main document text content or markdown"},
            "filename": {"type": "string", "description": "Output filename ending in .docx"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch Word"}
        },
        "required": ["title", "content"]
    }
)
def create_word_document(args: dict) -> str:
    """Create Word document via document_creator engine."""
    args_copy = dict(args) if isinstance(args, dict) else {}
    args_copy["format"] = "docx"
    if not args_copy.get("content"):
        args_copy["content"] = args_copy.get("title", "Document Content")
    return document_creator(args_copy)


@register_tool(
    name="create_pdf_document",
    description="Create a formatted PDF (.pdf) document and auto-launch in default PDF viewer.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main text content or markdown"},
            "filename": {"type": "string", "description": "Output filename ending in .pdf"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch PDF viewer"}
        },
        "required": ["title", "content"]
    }
)
def create_pdf_document(args: dict) -> str:
    """Create PDF document via document_creator engine."""
    args_copy = dict(args) if isinstance(args, dict) else {}
    args_copy["format"] = "pdf"
    if not args_copy.get("content"):
        args_copy["content"] = args_copy.get("title", "Document Content")
    return document_creator(args_copy)


@register_tool(
    name="generate_project_product_analysis",
    description="Generate a complete Product Analysis Report for B.R. JARVIS as Word (.docx) and PDF (.pdf) documents and auto-open them.",
    parameters={"type": "object", "properties": {}}
)
def generate_project_product_analysis(args: dict) -> str:
    """Generate complete Product Analysis report for B.R. JARVIS in Word & PDF formats."""
    doc_title = "Product Analysis Document: B.R. JARVIS"
    doc_text = """
# Product Analysis Document: B.R. JARVIS

## 1. Executive Summary
- **Product Name:** B.R. JARVIS (Advanced Agentic AI Operating System)
- **Identity:** Ultra-fast autonomous AI assistant designed for decisive multi-step reasoning, pair programming, and live desktop visual control.
- **Mission:** To eliminate conversational filler and deliver instant, high-precision software engineering actions.

> [!NOTE]
> B.R. JARVIS operates local-first with zero token latency optimizations and full execution autonomy.

## 2. Product Vision & Value Proposition
B.R. JARVIS shifts the paradigm from static autocomplete AI to an autonomous senior developer:
- **Zero-Filler Directive:** Delivers immediate, high-signal task resolution with minimal output tokens.
- **Local Gateway Integration:** Operates with unlimited request quotas via local gateway http://127.0.0.1:8045/v1.

## 3. Core Features & Capabilities
- **0-Token Intent Engine:** Executes common app launches, browser navigation, and diagnostics in 0ms with zero token cost.
- **Live OS Visual Controller:** Performs real-time desktop visual grounded automation (2160x1440 screen resolution).
- **Multi-Tab Excel & Document Generator:** Automatically creates formatted .xlsx, .docx, and .pdf analytical reports.

## 4. Architecture & Subsystems Inventory
| Subsystem | Components | Function |
| --- | --- | --- |
| Core Kernel | Native C FNV-1a bridge, EventBus runtime | Fast message routing |
| Action Suite | Live OS Control, Computer Control, RAG Library | Desktop & web automation |
| Tool Registry | 140+ registered tools | Plugin architecture |
| UI HUD | Tkinter glassmorphism display | Real-time status feedback |

## 5. Security & Execution Safety
- Local-first workspace execution with absolute path validation.
- AST compilation checks and security vulnerability scanning.
"""
    res_word = document_creator({"title": doc_title, "content": doc_text, "filename": "workspace/Reports/JARVIS_Product_Analysis.docx", "format": "docx", "auto_open": True})
    res_pdf = document_creator({"title": doc_title, "content": doc_text, "filename": "workspace/Reports/JARVIS_Product_Analysis.pdf", "format": "pdf", "auto_open": True})

    return f"{res_word}\n{res_pdf}"


@register_tool(
    name="generate_walkthrough",
    description="Generate a rich GitHub-flavored Markdown Walkthrough document (walkthrough.md) documenting technical changes, verification results, and file links.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title of the walkthrough"},
            "summary": {"type": "string", "description": "High-level summary of work accomplished"},
            "changes": {"type": "string", "description": "Detailed description or markdown list of changes made"},
            "verification": {"type": "string", "description": "Verification steps and automated test results"},
            "filename": {"type": "string", "description": "Target filename, default is walkthrough.md"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-open the generated walkthrough file"}
        },
        "required": ["title", "changes"]
    }
)
def generate_walkthrough(args: dict) -> str:
    """Generate a GitHub-flavored markdown walkthrough document."""
    title = args.get("title", "Task Walkthrough").strip()
    summary = args.get("summary", "").strip()
    changes = args.get("changes", "").strip()
    verification = args.get("verification", "").strip()
    filename = args.get("filename", "walkthrough.md").strip()
    auto_open = args.get("auto_open", False)

    content_lines = [f"# Walkthrough — {title}\n"]
    if summary:
        content_lines.append(f"{summary}\n")
    content_lines.append("## Changes Made\n")
    content_lines.append(f"{changes}\n")
    if verification:
        content_lines.append("## Verification Results\n")
        content_lines.append(f"{verification}\n")

    full_md = "\n".join(content_lines)
    out_path = _get_workspace_dir() / filename

    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(full_md, encoding="utf-8")
    except Exception as e:
        return f"Error writing walkthrough document: {e}"

    if auto_open and sys.platform == "win32":
        try:
            subprocess.Popen(["cmd", "/c", "start", "", str(out_path)], shell=False)
        except Exception:
            pass

    file_uri = out_path.as_uri()
    return f"⚡ Generated Walkthrough document successfully: [{filename}]({file_uri})"
