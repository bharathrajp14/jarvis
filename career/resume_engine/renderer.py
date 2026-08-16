# career/resume_engine/renderer.py — Unified HTML, DOCX, and PDF Resume Renderer for BR JARVIS
from __future__ import annotations

import logging
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from career.models import CareerProfile
from career.resume_engine.models import (
    ResumeSchema,
    SectionConfig,
    TemplateType,
    ThemeConfig,
)
from career.resume_engine.templates import TEMPLATES, get_template

logger = logging.getLogger("JARVIS.ResumeRenderer")

# Check DOCX & FPDF dependencies
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False


def _hex_to_rgb(hex_str: str) -> tuple[int, int, int]:
    """Convert hex color string (e.g. '#1B365D') to (R, G, B) integer tuple."""
    clean = hex_str.lstrip("#")
    if len(clean) == 3:
        clean = "".join(c * 2 for c in clean)
    if len(clean) != 6:
        return (0, 0, 0)
    try:
        return (int(clean[0:2], 16), int(clean[2:4], 16), int(clean[4:6], 16))
    except Exception:
        return (0, 0, 0)


class ResumeRenderer:
    """Master rendering engine supporting HTML, DOCX, and PDF with verified physical outputs."""

    # ── 1. Schema Generation from Master Profile ─────────────────────────────

    @classmethod
    def schema_from_profile(
        cls,
        profile: CareerProfile,
        target_role: Optional[str] = None,
        template_id: Union[TemplateType, str] = TemplateType.ATS_CLASSIC,
        theme: Optional[ThemeConfig] = None,
    ) -> ResumeSchema:
        """Construct a rendered view schema snapshot from canonical CareerProfile."""
        tmpl_def = get_template(template_id)
        active_theme = theme or tmpl_def.default_theme

        links = {}
        if profile.contact.linkedin_url:
            links["LinkedIn"] = profile.contact.linkedin_url
        if profile.contact.github_url:
            links["GitHub"] = profile.contact.github_url
        if profile.contact.portfolio_url:
            links["Portfolio"] = profile.contact.portfolio_url

        role = target_role or (profile.preferences.target_roles[0] if profile.preferences.target_roles else "Software Engineer")

        return ResumeSchema(
            title=f"{profile.contact.full_name} — {role}",
            target_role=role,
            template_id=tmpl_def.template_id,
            theme=active_theme,
            sections=tmpl_def.default_sections,
            full_name=profile.contact.full_name,
            contact_email=profile.contact.email,
            contact_phone=profile.contact.phone,
            location=profile.contact.location,
            links=links,
            summary=profile.summary,
            experience=[e.to_dict() for e in profile.experience],
            education=[e.to_dict() for e in profile.education],
            skills=[s.to_dict() for s in profile.skills],
            projects=[p.to_dict() for p in profile.projects],
            certifications=[c.to_dict() for c in profile.certifications],
            achievements=[a.to_dict() for a in profile.achievements],
        )

    # ── 2. HTML Rendering ───────────────────────────────────────────────────

    @classmethod
    def render_html(cls, resume: ResumeSchema) -> str:
        """Render standalone, responsive, print-ready HTML resume with typography & theme tokens."""
        tmpl_def = get_template(resume.template_id)
        theme = resume.theme

        # Links line
        links_parts = []
        if resume.contact_email:
            links_parts.append(f'<a href="mailto:{resume.contact_email}">{resume.contact_email}</a>')
        if resume.contact_phone:
            links_parts.append(f'<span>{resume.contact_phone}</span>')
        if resume.location:
            links_parts.append(f'<span>{resume.location}</span>')
        for label, url in resume.links.items():
            clean_url = url.replace("https://", "").replace("http://", "")
            links_parts.append(f'<a href="{url}" target="_blank" rel="noopener">{label}: {clean_url}</a>')
        contact_line_html = " • ".join(links_parts)

        # Build sections by configured order
        sorted_sections = sorted([s for s in resume.sections if s.visible], key=lambda x: x.order)
        sections_html = []

        for sec in sorted_sections:
            sec_id = sec.section_id.lower()

            # Summary
            if sec_id == "summary" and resume.summary:
                sections_html.append(f"""
                    <div class="resume-section" id="section-summary">
                        <div class="section-title">{sec.title}</div>
                        <p class="summary-text">{resume.summary}</p>
                    </div>
                """)

            # Skills
            elif sec_id == "skills" and resume.skills:
                skills_rows = []
                for sc in resume.skills:
                    cat_name = sc.get("category", "Skills")
                    skills_list = sc.get("skills", [])
                    if skills_list:
                        pill_tags = "".join(f'<span class="skill-tag">{s}</span>' for s in skills_list)
                        skills_rows.append(f"""
                            <div class="skill-category-row">
                                <span class="skill-category-name">{cat_name}:</span>
                                <span class="skill-tags-group">{pill_tags}</span>
                            </div>
                        """)
                if skills_rows:
                    sections_html.append(f"""
                        <div class="resume-section" id="section-skills">
                            <div class="section-title">{sec.title}</div>
                            <div class="skills-container">{"".join(skills_rows)}</div>
                        </div>
                    """)

            # Experience
            elif sec_id == "experience" and resume.experience:
                exp_entries = []
                for exp in resume.experience:
                    co = exp.get("company", "")
                    title = exp.get("title", "")
                    loc = exp.get("location", "")
                    s_date = exp.get("start_date", "")
                    e_date = exp.get("end_date", "Present")
                    dates = f"{s_date} – {e_date}" if s_date else e_date
                    bullets = "".join(f'<li>{r}</li>' for r in exp.get("responsibilities", []) + exp.get("achievements", []))
                    
                    tech_html = ""
                    if exp.get("technologies"):
                        tech_html = f'<div class="entry-technologies"><strong>Tech Stack:</strong> {", ".join(exp.get("technologies"))}</div>'

                    exp_entries.append(f"""
                        <div class="job-entry">
                            <div class="job-header">
                                <span class="job-title">{title}</span>
                                <span class="job-dates">{dates}</span>
                            </div>
                            <div class="job-sub">
                                <span class="job-company">{co}</span>
                                <span class="job-location">{loc}</span>
                            </div>
                            <ul class="bullet-list">{bullets}</ul>
                            {tech_html}
                        </div>
                    """)
                if exp_entries:
                    sections_html.append(f"""
                        <div class="resume-section" id="section-experience">
                            <div class="section-title">{sec.title}</div>
                            {"".join(exp_entries)}
                        </div>
                    """)

            # Projects
            elif sec_id == "projects" and resume.projects:
                proj_entries = []
                for p in resume.projects:
                    name = p.get("name", "")
                    role = p.get("role", "")
                    desc = p.get("description", "")
                    tech = ", ".join(p.get("technologies", []))
                    bullets = "".join(f'<li>{h}</li>' for h in p.get("highlights", []))
                    url_html = f'<a href="{p.get("url")}" target="_blank" class="proj-link">{p.get("url")}</a>' if p.get("url") else ""

                    proj_entries.append(f"""
                        <div class="project-entry">
                            <div class="job-header">
                                <span class="job-title">{name} {f"— <em>{role}</em>" if role else ""}</span>
                                <span class="job-dates">{url_html}</span>
                            </div>
                            {f'<p class="proj-desc">{desc}</p>' if desc else ''}
                            {f'<ul class="bullet-list">{bullets}</ul>' if bullets else ''}
                            {f'<div class="entry-technologies"><strong>Technologies:</strong> {tech}</div>' if tech else ''}
                        </div>
                    """)
                if proj_entries:
                    sections_html.append(f"""
                        <div class="resume-section" id="section-projects">
                            <div class="section-title">{sec.title}</div>
                            {"".join(proj_entries)}
                        </div>
                    """)

            # Education
            elif sec_id == "education" and resume.education:
                edu_entries = []
                for edu in resume.education:
                    inst = edu.get("institution", "")
                    deg = edu.get("degree", "")
                    fld = edu.get("field_of_study", "")
                    full_deg = f"{deg} in {fld}" if fld else deg
                    dates = f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
                    bullets = "".join(f'<li>{h}</li>' for h in edu.get("highlights", []))

                    edu_entries.append(f"""
                        <div class="edu-entry">
                            <div class="job-header">
                                <span class="job-title">{full_deg}</span>
                                <span class="job-dates">{dates}</span>
                            </div>
                            <div class="job-sub">
                                <span class="job-company">{inst}</span>
                                <span class="job-location">{edu.get('location', '')}</span>
                            </div>
                            {f'<ul class="bullet-list">{bullets}</ul>' if bullets else ''}
                        </div>
                    """)
                if edu_entries:
                    sections_html.append(f"""
                        <div class="resume-section" id="section-education">
                            <div class="section-title">{sec.title}</div>
                            {"".join(edu_entries)}
                        </div>
                    """)

            # Certifications
            elif sec_id == "certifications" and resume.certifications:
                cert_items = []
                for c in resume.certifications:
                    name = c.get("name", "")
                    issuer = c.get("issuer", "")
                    date = c.get("issue_date", "")
                    cert_items.append(f'<li><strong>{name}</strong> — {issuer} {f"({date})" if date else ""}</li>')
                if cert_items:
                    sections_html.append(f"""
                        <div class="resume-section" id="section-certifications">
                            <div class="section-title">{sec.title}</div>
                            <ul class="bullet-list">{"".join(cert_items)}</ul>
                        </div>
                    """)

            # Achievements
            elif sec_id == "achievements" and resume.achievements:
                ach_items = []
                for a in resume.achievements:
                    title = a.get("title", "")
                    desc = a.get("description", "")
                    date = a.get("date", "")
                    ach_items.append(f'<li><strong>{title}</strong>{f" ({date})" if date else ""}: {desc}</li>')
                if ach_items:
                    sections_html.append(f"""
                        <div class="resume-section" id="section-achievements">
                            <div class="section-title">{sec.title}</div>
                            <ul class="bullet-list">{"".join(ach_items)}</ul>
                        </div>
                    """)

        html_body = "".join(sections_html)

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{resume.title}</title>
    <link href="https://fonts.googleapis.com/css2?family=Cinzel:wght@600;700&family=Inter:wght@300;400;500;600;700;800&family=Outfit:wght@400;600;700&family=Share+Tech+Mono&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: {theme.primary_color};
            --secondary: {theme.secondary_color};
            --accent: {theme.accent_color};
            --text: {theme.text_color};
            --muted: {theme.muted_text};
            --bg: {theme.background_color};
            --card-bg: {theme.card_background};
            --font-heading: {theme.font_heading};
            --font-body: {theme.font_body};
            --font-code: {theme.font_code};
            --font-size: {theme.font_size_base};
            --line-height: {theme.line_height};
        }}
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: var(--font-body);
            font-size: var(--font-size);
            line-height: var(--line-height);
            color: var(--text);
            background: #E2E8F0;
            padding: 24px;
        }}
        .resume-page {{
            max-width: 820px;
            margin: 0 auto;
            background: var(--bg);
            padding: {theme.margin_pt}pt;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
            border-radius: {theme.border_radius};
        }}
        a {{ color: var(--accent); text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        .summary-text {{ font-size: 9.5pt; line-height: 1.5; color: var(--text); }}
        .skill-category-row {{ margin-bottom: 4px; font-size: 9.5pt; }}
        .skill-category-name {{ font-weight: 600; color: var(--primary); margin-right: 6px; }}
        .skill-tag {{ display: inline-block; background: var(--card-bg); color: var(--text); padding: 1px 6px; border-radius: 3px; font-size: 8.5pt; margin: 1px 3px 1px 0; border: 1px solid #E2E8F0; }}
        .entry-technologies {{ font-size: 8.5pt; color: var(--muted); margin-top: 3px; }}
        .proj-desc {{ font-size: 9pt; margin-top: 2px; color: var(--text); }}
        
        {tmpl_def.css_template}

        @media print {{
            body {{ background: #FFFFFF; padding: 0; }}
            .resume-page {{ box-shadow: none; max-width: 100%; border-radius: 0; padding: {theme.margin_pt}pt; }}
            @page {{ margin: 0; size: letter portrait; }}
        }}
    </style>
</head>
<body>
    <div class="resume-page">
        <header class="resume-header">
            <h1 class="candidate-name">{resume.full_name}</h1>
            <div class="candidate-title">{resume.target_role}</div>
            <div class="contact-line">{contact_line_html}</div>
        </header>
        <main class="resume-body">
            {html_body}
        </main>
    </div>
</body>
</html>"""

    # ── 3. DOCX Rendering ───────────────────────────────────────────────────

    @classmethod
    def render_docx(cls, resume: ResumeSchema, output_path: Union[str, Path]) -> bool:
        """Generate high-fidelity Microsoft Word DOCX document with typography & styling."""
        if not _DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

        out_file = Path(output_path).resolve()
        out_file.parent.mkdir(parents=True, exist_ok=True)

        doc = docx.Document()
        theme = resume.theme
        p_rgb = _hex_to_rgb(theme.primary_color)
        s_rgb = _hex_to_rgb(theme.secondary_color)
        a_rgb = _hex_to_rgb(theme.accent_color)
        m_rgb = _hex_to_rgb(theme.muted_text)

        # Page margins
        for section in doc.sections:
            margin_in = Inches(theme.margin_pt / 72.0)
            section.top_margin = margin_in
            section.bottom_margin = margin_in
            section.left_margin = margin_in
            section.right_margin = margin_in

        # Header: Name
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if resume.template_id in (TemplateType.EXECUTIVE, TemplateType.ATS_CLASSIC, TemplateType.FRESH_GRADUATE) else WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_p.add_run(resume.full_name)
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(*p_rgb)
        name_p.paragraph_format.space_after = Pt(2)

        # Title
        if resume.target_role:
            title_p = doc.add_paragraph()
            title_p.alignment = name_p.alignment
            title_run = title_p.add_run(resume.target_role.upper())
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(*a_rgb)
            title_p.paragraph_format.space_after = Pt(4)

        # Contact Line
        contact_parts = [p for p in [resume.contact_email, resume.contact_phone, resume.location] if p]
        for lbl, url in resume.links.items():
            contact_parts.append(f"{lbl}: {url.replace('https://', '')}")
        
        contact_p = doc.add_paragraph()
        contact_p.alignment = name_p.alignment
        contact_run = contact_p.add_run("  •  ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(*m_rgb)
        contact_p.paragraph_format.space_after = Pt(14)

        def add_section_header(title: str):
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(12)
            sp.paragraph_format.space_after = Pt(4)
            s_run = sp.add_run(title.upper())
            s_run.font.size = Pt(11)
            s_run.font.bold = True
            s_run.font.color.rgb = RGBColor(*p_rgb)

        # Sections
        sorted_sections = sorted([s for s in resume.sections if s.visible], key=lambda x: x.order)

        for sec in sorted_sections:
            sec_id = sec.section_id.lower()

            if sec_id == "summary" and resume.summary:
                add_section_header(sec.title)
                sum_p = doc.add_paragraph()
                sum_run = sum_p.add_run(resume.summary)
                sum_run.font.size = Pt(9.5)
                sum_p.paragraph_format.space_after = Pt(6)

            elif sec_id == "skills" and resume.skills:
                add_section_header(sec.title)
                for sc in resume.skills:
                    c_name = sc.get("category", "Technical Skills")
                    s_list = ", ".join(sc.get("skills", []))
                    if s_list:
                        sk_p = doc.add_paragraph()
                        cat_run = sk_p.add_run(f"{c_name}: ")
                        cat_run.font.bold = True
                        cat_run.font.size = Pt(9.5)
                        cat_run.font.color.rgb = RGBColor(*s_rgb)
                        val_run = sk_p.add_run(s_list)
                        val_run.font.size = Pt(9.5)
                        sk_p.paragraph_format.space_after = Pt(2)

            elif sec_id == "experience" and resume.experience:
                add_section_header(sec.title)
                for exp in resume.experience:
                    jp = doc.add_paragraph()
                    jp.paragraph_format.space_before = Pt(6)
                    jp.paragraph_format.space_after = Pt(1)
                    
                    t_run = jp.add_run(exp.get("title", ""))
                    t_run.font.bold = True
                    t_run.font.size = Pt(10)
                    t_run.font.color.rgb = RGBColor(*p_rgb)

                    dates = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
                    d_run = jp.add_run(f" | {dates}")
                    d_run.font.size = Pt(9)
                    d_run.font.italic = True
                    d_run.font.color.rgb = RGBColor(*m_rgb)

                    co_p = doc.add_paragraph()
                    co_p.paragraph_format.space_after = Pt(3)
                    co_run = co_p.add_run(f"{exp.get('company', '')} — {exp.get('location', '')}")
                    co_run.font.bold = True
                    co_run.font.size = Pt(9.5)
                    co_run.font.color.rgb = RGBColor(*s_rgb)

                    for b in exp.get("responsibilities", []) + exp.get("achievements", []):
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.paragraph_format.space_after = Pt(1.5)
                        b_run = bp.add_run(b)
                        b_run.font.size = Pt(9)

            elif sec_id == "projects" and resume.projects:
                add_section_header(sec.title)
                for p in resume.projects:
                    pp = doc.add_paragraph()
                    pp.paragraph_format.space_before = Pt(5)
                    pp.paragraph_format.space_after = Pt(2)
                    p_name = pp.add_run(p.get("name", ""))
                    p_name.font.bold = True
                    p_name.font.size = Pt(10)
                    p_name.font.color.rgb = RGBColor(*p_rgb)
                    if p.get("role"):
                        r_run = pp.add_run(f" ({p.get('role')})")
                        r_run.font.italic = True
                        r_run.font.size = Pt(9)

                    if p.get("description"):
                        dp = doc.add_paragraph()
                        dp.paragraph_format.space_after = Pt(2)
                        d_run = dp.add_run(p.get("description"))
                        d_run.font.size = Pt(9)

                    for h in p.get("highlights", []):
                        hp = doc.add_paragraph(style="List Bullet")
                        hp.paragraph_format.space_after = Pt(1)
                        h_run = hp.add_run(h)
                        h_run.font.size = Pt(9)

            elif sec_id == "education" and resume.education:
                add_section_header(sec.title)
                for edu in resume.education:
                    ep = doc.add_paragraph()
                    ep.paragraph_format.space_before = Pt(4)
                    ep.paragraph_format.space_after = Pt(1)
                    deg_run = ep.add_run(f"{edu.get('degree', '')} in {edu.get('field_of_study', '')}")
                    deg_run.font.bold = True
                    deg_run.font.size = Pt(9.5)
                    deg_run.font.color.rgb = RGBColor(*p_rgb)

                    inst_p = doc.add_paragraph()
                    inst_p.paragraph_format.space_after = Pt(2)
                    inst_run = inst_p.add_run(f"{edu.get('institution', '')} ({edu.get('start_date', '')}–{edu.get('end_date', '')})")
                    inst_run.font.size = Pt(9)
                    inst_run.font.color.rgb = RGBColor(*m_rgb)

            elif sec_id == "certifications" and resume.certifications:
                add_section_header(sec.title)
                for c in resume.certifications:
                    cp = doc.add_paragraph(style="List Bullet")
                    cp.paragraph_format.space_after = Pt(1.5)
                    c_run = cp.add_run(f"{c.get('name')} — {c.get('issuer', '')}")
                    c_run.font.size = Pt(9)

        doc.save(str(out_file))
        logger.info(f"📄 DOCX Resume written: {out_file} ({out_file.stat().st_size:,} bytes)")
        return True

    # ── 4. PDF Rendering ────────────────────────────────────────────────────

    @classmethod
    def render_pdf(cls, resume: ResumeSchema, output_path: Union[str, Path]) -> bool:
        """Generate vector PDF resume using fpdf2 or Playwright fallback."""
        if not _FPDF_AVAILABLE:
            # Render HTML then write to disk
            html_content = cls.render_html(resume)
            html_temp = Path(output_path).with_suffix(".html")
            html_temp.write_text(html_content, encoding="utf-8")
            raise RuntimeError("fpdf2 is not installed. Install with: pip install fpdf2")

        out_file = Path(output_path).resolve()
        out_file.parent.mkdir(parents=True, exist_ok=True)

        theme = resume.theme
        p_rgb = _hex_to_rgb(theme.primary_color)
        s_rgb = _hex_to_rgb(theme.secondary_color)
        a_rgb = _hex_to_rgb(theme.accent_color)
        m_rgb = _hex_to_rgb(theme.muted_text)
        t_rgb = _hex_to_rgb(theme.text_color)

        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.set_auto_page_break(auto=True, margin=12)
        pdf.add_page()
        pdf.set_margins(12, 12, 12)

        # Header: Name
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(*p_rgb)
        pdf.cell(0, 8, resume.full_name.encode('latin-1', 'replace').decode('latin-1'), align="C", new_x="LMARGIN", new_y="NEXT")

        # Target Role
        if resume.target_role:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*a_rgb)
            pdf.cell(0, 5, resume.target_role.upper().encode('latin-1', 'replace').decode('latin-1'), align="C", new_x="LMARGIN", new_y="NEXT")

        # Contact Line
        contact_parts = [p for p in [resume.contact_email, resume.contact_phone, resume.location] if p]
        for lbl, url in resume.links.items():
            contact_parts.append(f"{lbl}: {url.replace('https://', '')}")
        
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*m_rgb)
        contact_text = "  |  ".join(contact_parts)
        pdf.cell(0, 5, contact_text.encode('latin-1', 'replace').decode('latin-1'), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        def add_pdf_section(title: str):
            pdf.ln(3)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*p_rgb)
            pdf.cell(pdf.epw, 5, title.upper().encode('latin-1', 'replace').decode('latin-1'), align="L")
            pdf.ln(5)
            # Horizontal rule
            pdf.set_draw_color(*p_rgb)
            pdf.set_line_width(0.3)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + pdf.epw, pdf.get_y())
            pdf.ln(2)
            pdf.set_x(pdf.l_margin)

        sorted_sections = sorted([s for s in resume.sections if s.visible], key=lambda x: x.order)

        for sec in sorted_sections:
            sec_id = sec.section_id.lower()

            if sec_id == "summary" and resume.summary:
                add_pdf_section(sec.title)
                pdf.set_x(pdf.l_margin)
                pdf.set_font("Helvetica", "", 8.5)
                pdf.set_text_color(*t_rgb)
                pdf.multi_cell(pdf.epw, 4.2, resume.summary.encode('latin-1', 'replace').decode('latin-1'))
                pdf.ln(1)

            elif sec_id == "skills" and resume.skills:
                add_pdf_section(sec.title)
                for sc in resume.skills:
                    c_name = sc.get("category", "Skills")
                    s_list = ", ".join(sc.get("skills", []))
                    if s_list:
                        pdf.set_x(pdf.l_margin)
                        pdf.set_font("Helvetica", "B", 8.5)
                        pdf.set_text_color(*s_rgb)
                        clean_line = f"{c_name}: {s_list}".encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(pdf.epw, 4.0, clean_line)
                pdf.ln(1)

            elif sec_id == "experience" and resume.experience:
                add_pdf_section(sec.title)
                for exp in resume.experience:
                    title = exp.get("title", "")
                    co = exp.get("company", "")
                    dates = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
                    
                    pdf.set_x(pdf.l_margin)
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(*p_rgb)
                    header_line = f"{title} | {dates}".encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(pdf.epw, 4.5, header_line, align="L")
                    pdf.ln(4.5)

                    pdf.set_x(pdf.l_margin)
                    pdf.set_font("Helvetica", "B", 8.5)
                    pdf.set_text_color(*s_rgb)
                    pdf.cell(pdf.epw, 4, f"{co} — {exp.get('location', '')}".encode('latin-1', 'replace').decode('latin-1'), align="L")
                    pdf.ln(4)

                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_text_color(*t_rgb)
                    for b in exp.get("responsibilities", []) + exp.get("achievements", []):
                        pdf.set_x(pdf.l_margin)
                        clean_b = f"  -  {b}".encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(pdf.epw, 3.8, clean_b)
                    pdf.ln(1.5)

            elif sec_id == "projects" and resume.projects:
                add_pdf_section(sec.title)
                for p in resume.projects:
                    pdf.set_x(pdf.l_margin)
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(*p_rgb)
                    role_str = f" ({p.get('role')})" if p.get("role") else ""
                    p_title = f"{p.get('name', '')}{role_str}"
                    pdf.cell(pdf.epw, 4.5, p_title.encode('latin-1', 'replace').decode('latin-1'), align="L")
                    pdf.ln(4.5)

                    if p.get("description"):
                        pdf.set_x(pdf.l_margin)
                        pdf.set_font("Helvetica", "", 8)
                        pdf.set_text_color(*t_rgb)
                        pdf.multi_cell(pdf.epw, 3.8, p.get("description", "").encode('latin-1', 'replace').decode('latin-1'))

                    for h in p.get("highlights", []):
                        pdf.set_x(pdf.l_margin)
                        pdf.set_font("Helvetica", "", 8)
                        pdf.set_text_color(*t_rgb)
                        clean_h = f"  -  {h}".encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(pdf.epw, 3.8, clean_h)
                    pdf.ln(1)

            elif sec_id == "education" and resume.education:
                add_pdf_section(sec.title)
                for edu in resume.education:
                    deg = f"{edu.get('degree', '')} in {edu.get('field_of_study', '')}"
                    dates = f"{edu.get('start_date', '')}–{edu.get('end_date', '')}"
                    
                    pdf.set_x(pdf.l_margin)
                    pdf.set_font("Helvetica", "B", 8.5)
                    pdf.set_text_color(*p_rgb)
                    pdf.cell(pdf.epw * 0.75, 4, deg.encode('latin-1', 'replace').decode('latin-1'))
                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_text_color(*m_rgb)
                    pdf.cell(pdf.epw * 0.25, 4, dates.encode('latin-1', 'replace').decode('latin-1'), align="R")
                    pdf.ln(4)

                    pdf.set_x(pdf.l_margin)
                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_text_color(*s_rgb)
                    pdf.cell(pdf.epw, 3.8, edu.get("institution", "").encode('latin-1', 'replace').decode('latin-1'))
                    pdf.ln(4)

            elif sec_id == "certifications" and resume.certifications:
                add_pdf_section(sec.title)
                pdf.set_font("Helvetica", "", 8)
                pdf.set_text_color(*t_rgb)
                for c in resume.certifications:
                    pdf.set_x(pdf.l_margin)
                    line = f"  -  {c.get('name')} — {c.get('issuer', '')} ({c.get('issue_date', '')})"
                    pdf.multi_cell(pdf.epw, 3.8, line.encode('latin-1', 'replace').decode('latin-1'))
                pdf.ln(1)

        pdf.output(str(out_file))
        logger.info(f"📄 PDF Resume written: {out_file} ({out_file.stat().st_size:,} bytes)")
        return True
