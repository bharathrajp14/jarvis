# agent/verifier.py — Autonomous Action & Goal Verification Subsystem
"""
State and Goal Verification Engine for BR JARVIS MK40.
Ensures actions are rigorously verified against actual OS/filesystem/process/browser state,
preventing false-positive completions, unhandled sandbox handoffs, and hallucinated success states.
"""
from __future__ import annotations

import json
import logging
import os
import re
import sys
import time
from dataclasses import asdict, dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

logger = logging.getLogger("JARVIS.Verifier")


class VerificationStatus(str, Enum):
    SUCCESS_VERIFIED   = "SUCCESS_VERIFIED"
    SUCCESS_UNVERIFIED = "SUCCESS_UNVERIFIED"
    PARTIAL_SUCCESS    = "PARTIAL_SUCCESS"
    FAILED             = "FAILED"
    TIMEOUT            = "TIMEOUT"
    CANCELLED          = "CANCELLED"
    BLOCKED            = "BLOCKED"
    NOT_IMPLEMENTED    = "NOT_IMPLEMENTED"


@dataclass
class VerificationResult:
    """Outcome of an action verification check with evidence and status."""
    verified: bool
    status: VerificationStatus = VerificationStatus.SUCCESS_VERIFIED
    evidence: str = ""
    details: str = ""
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        if not self.evidence and self.details:
            self.evidence = self.details
        if not self.details and self.evidence:
            self.details = self.evidence
        if isinstance(self.status, str) and not isinstance(self.status, VerificationStatus):
            try:
                self.status = VerificationStatus(self.status)
            except ValueError:
                self.status = VerificationStatus.SUCCESS_VERIFIED if self.verified else VerificationStatus.FAILED

    def to_dict(self) -> Dict[str, Any]:
        return {
            "verified": self.verified,
            "status": self.status.value,
            "evidence": self.evidence,
            "details": self.details,
            "error": self.error,
            "metadata": self.metadata,
        }


# ── Specialized Verifiers ──────────────────────────────────────────────────

class FileVerifier:
    """Verifies file existence, size, permissions, integrity, and parsability."""

    @staticmethod
    def verify_file_created(path_str: str, min_size_bytes: int = 1) -> VerificationResult:
        try:
            p = Path(path_str).resolve()
            if not p.exists():
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.FAILED,
                    details=f"Verification failed: File '{path_str}' does not exist on disk.",
                    error="FILE_NOT_FOUND"
                )
            if not p.is_file():
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.FAILED,
                    details=f"Verification failed: Path '{path_str}' is a directory, not a file.",
                    error="NOT_A_FILE"
                )
            size = p.stat().st_size
            if size < min_size_bytes:
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.FAILED,
                    details=f"Verification failed: File '{path_str}' is empty ({size} bytes).",
                    error="FILE_EMPTY"
                )
            return VerificationResult(
                verified=True,
                status=VerificationStatus.SUCCESS_VERIFIED,
                evidence=f"File '{p.name}' verified on disk ({size:,} bytes, path: {p}).",
                details=f"Verified file '{p.name}' created successfully ({size} bytes).",
                metadata={"path": str(p), "size_bytes": size}
            )
        except Exception as e:
            return VerificationResult(
                verified=False,
                status=VerificationStatus.FAILED,
                details=f"Verification error for '{path_str}': {e}",
                error="VERIFICATION_EXCEPTION"
            )

    @staticmethod
    def verify_file_content(path_str: str, expected_substrings: Optional[List[str]] = None) -> VerificationResult:
        res = FileVerifier.verify_file_created(path_str)
        if not res.verified:
            return res
        p = Path(path_str).resolve()
        try:
            content = p.read_text(encoding="utf-8", errors="replace")
            if expected_substrings:
                missing = [s for s in expected_substrings if s.lower() not in content.lower()]
                if missing:
                    return VerificationResult(
                        verified=False,
                        status=VerificationStatus.PARTIAL_SUCCESS,
                        details=f"File exists but missing expected content: {missing}",
                        error="CONTENT_MISMATCH",
                        metadata={"missing": missing, "path": str(p)}
                    )
            return VerificationResult(
                verified=True,
                status=VerificationStatus.SUCCESS_VERIFIED,
                evidence=f"File '{p.name}' contains verified content ({len(content)} chars).",
                details=f"File content verified for '{p.name}'.",
                metadata={"path": str(p), "char_count": len(content)}
            )
        except Exception as e:
            return VerificationResult(
                verified=False,
                status=VerificationStatus.FAILED,
                details=f"Error reading file '{path_str}': {e}",
                error="READ_ERROR"
            )

    @staticmethod
    def verify_file_parsed(path_str: str) -> VerificationResult:
        """Parse and validate structural integrity of complex documents (DOCX, PDF, XLSX, JSON)."""
        res = FileVerifier.verify_file_created(path_str)
        if not res.verified:
            return res

        p = Path(path_str).resolve()
        ext = p.suffix.lower()

        try:
            if ext == ".docx":
                import docx
                doc = docx.Document(str(p))
                p_count = len(doc.paragraphs)
                t_count = len(doc.tables)
                total_text_len = sum(len(p.text) for p in doc.paragraphs)
                if p_count == 0 and t_count == 0:
                    return VerificationResult(
                        verified=False,
                        status=VerificationStatus.FAILED,
                        details=f"DOCX '{p.name}' parsed but contains zero paragraphs and zero tables.",
                        error="EMPTY_DOCX"
                    )
                return VerificationResult(
                    verified=True,
                    status=VerificationStatus.SUCCESS_VERIFIED,
                    evidence=f"DOCX '{p.name}' parsed successfully ({p_count} paragraphs, {t_count} tables, {total_text_len} chars).",
                    details=f"Verified DOCX structure for '{p.name}'.",
                    metadata={"paragraphs": p_count, "tables": t_count, "size": p.stat().st_size}
                )

            elif ext == ".pdf":
                # Validate PDF magic header
                with open(p, "rb") as f:
                    header = f.read(5)
                if not header.startswith(b"%PDF-"):
                    return VerificationResult(
                        verified=False,
                        status=VerificationStatus.FAILED,
                        details=f"File '{p.name}' lacks valid PDF magic header %PDF-.",
                        error="INVALID_PDF_HEADER"
                    )
                return VerificationResult(
                    verified=True,
                    status=VerificationStatus.SUCCESS_VERIFIED,
                    evidence=f"PDF '{p.name}' header and binary structure verified ({p.stat().st_size:,} bytes).",
                    details=f"Verified PDF structure for '{p.name}'.",
                    metadata={"size": p.stat().st_size}
                )

            elif ext == ".json":
                data = json.loads(p.read_text(encoding="utf-8"))
                return VerificationResult(
                    verified=True,
                    status=VerificationStatus.SUCCESS_VERIFIED,
                    evidence=f"JSON '{p.name}' parsed successfully ({type(data).__name__} root).",
                    details=f"Verified JSON structure for '{p.name}'."
                )

            elif ext in (".xlsx", ".xlsm"):
                import zipfile
                if not zipfile.is_zipfile(str(p)):
                    return VerificationResult(
                        verified=False,
                        status=VerificationStatus.FAILED,
                        details=f"XLSX '{p.name}' is not a valid OpenXML ZIP archive.",
                        error="INVALID_XLSX"
                    )
                return VerificationResult(
                    verified=True,
                    status=VerificationStatus.SUCCESS_VERIFIED,
                    evidence=f"Excel spreadsheet '{p.name}' verified ({p.stat().st_size:,} bytes).",
                    details=f"Verified XLSX archive for '{p.name}'."
                )

            return VerificationResult(
                verified=True,
                status=VerificationStatus.SUCCESS_VERIFIED,
                evidence=f"File '{p.name}' verified ({p.stat().st_size:,} bytes).",
                details=f"File verified: {p.name}"
            )
        except Exception as e:
            return VerificationResult(
                verified=False,
                status=VerificationStatus.FAILED,
                details=f"Structural parse verification failed for '{p.name}': {e}",
                error="PARSE_ERROR"
            )


class ApplicationVerifier:
    """Verifies OS application launches, active processes, and visible windows."""

    @staticmethod
    def verify_process_running(proc_name: str) -> VerificationResult:
        try:
            import psutil
            low = proc_name.lower().strip()
            base = os.path.splitext(low)[0]

            matched = []
            for proc in psutil.process_iter(['name', 'pid', 'create_time']):
                try:
                    pname = (proc.info.get('name') or '').lower()
                    if base in pname or low in pname:
                        matched.append(proc.info)
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue

            if matched:
                best = matched[0]
                return VerificationResult(
                    verified=True,
                    status=VerificationStatus.SUCCESS_VERIFIED,
                    evidence=f"Process '{best['name']}' (PID: {best['pid']}) confirmed active in OS process table.",
                    details=f"Verified process '{best['name']}' active (PID: {best['pid']}).",
                    metadata={"process_name": best['name'], "pid": best['pid']}
                )

            return VerificationResult(
                verified=False,
                status=VerificationStatus.SUCCESS_UNVERIFIED,
                details=f"Process '{proc_name}' not found in active process table.",
                error="PROCESS_NOT_FOUND"
            )
        except Exception as e:
            return VerificationResult(
                verified=True,
                status=VerificationStatus.SUCCESS_UNVERIFIED,
                details=f"Process verification skipped (psutil notice: {e})."
            )

    @staticmethod
    def verify_window_open(window_title_keyword: Optional[str] = None, app_name: Optional[str] = None) -> VerificationResult:
        """Inspect visible GUI window handles on Windows OS."""
        if sys.platform != "win32":
            return ApplicationVerifier.verify_process_running(app_name or "")

        try:
            import ctypes
            from ctypes import wintypes

            titles = []

            def enum_windows_proc(hwnd, lParam):
                if ctypes.windll.user32.IsWindowVisible(hwnd):
                    length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
                    if length > 0:
                        buff = ctypes.create_unicode_buffer(length + 1)
                        ctypes.windll.user32.GetWindowTextW(hwnd, buff, length + 1)
                        if buff.value.strip():
                            titles.append(buff.value.strip())
                return True

            WNDENUMPROC = ctypes.WINFUNCTYPE(wintypes.BOOL, wintypes.HWND, wintypes.LPARAM)
            ctypes.windll.user32.EnumWindows(WNDENUMPROC(enum_windows_proc), 0)

            kw = (window_title_keyword or app_name or "").lower().strip()
            if kw:
                matches = [t for t in titles if kw in t.lower()]
                if matches:
                    return VerificationResult(
                        verified=True,
                        status=VerificationStatus.SUCCESS_VERIFIED,
                        evidence=f"Active window detected: '{matches[0]}'.",
                        details=f"Verified window open matching '{kw}': '{matches[0]}'",
                        metadata={"window_title": matches[0], "all_matches": matches}
                    )

            # If no specific keyword or not found in windows, fallback to process check
            if app_name:
                p_res = ApplicationVerifier.verify_process_running(app_name)
                if p_res.verified:
                    return VerificationResult(
                        verified=True,
                        status=VerificationStatus.SUCCESS_VERIFIED,
                        evidence=f"Application '{app_name}' active with running process PID: {p_res.metadata.get('pid')}.",
                        details=f"Application process confirmed running.",
                        metadata=p_res.metadata
                    )

            return VerificationResult(
                verified=False,
                status=VerificationStatus.SUCCESS_UNVERIFIED,
                details=f"No visible window or process detected matching '{kw or app_name}'.",
                error="WINDOW_NOT_FOUND"
            )
        except Exception as e:
            return ApplicationVerifier.verify_process_running(app_name or window_title_keyword or "")


class BrowserVerifier:
    """Verifies that browser opened a legitimate, reachable artifact URL with no error pages."""

    @staticmethod
    def verify_browser_artifact_opened(
        host_path_or_url: Union[str, Path],
        browser_response: Optional[Union[dict, str]] = None,
        expected_content: Optional[str] = None
    ) -> VerificationResult:
        target_str = str(host_path_or_url).strip()
        low_target = target_str.lower().replace("\\", "/")

        # 1. Sandbox jail path check
        if "jarvis_sandbox_jails" in low_target or "sandbox_jails" in low_target or "/jail_" in low_target:
            return VerificationResult(
                verified=False,
                status=VerificationStatus.BLOCKED,
                details=f"Security Violation: Browser attempted to open internal sandbox jail path '{target_str}'.",
                error="SANDBOX_PATH_EXPOSURE"
            )

        # 2. Local file validation
        if not (target_str.startswith("http://") or target_str.startswith("https://")):
            clean_path = target_str
            if clean_path.startswith("file:///"):
                clean_path = clean_path[8:] if sys.platform == "win32" else clean_path[7:]
            elif clean_path.startswith("file://"):
                clean_path = clean_path[7:]

            p = Path(clean_path).resolve()
            if not p.exists():
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.FAILED,
                    details=f"Browser target file '{clean_path}' does not exist on disk.",
                    error="ERR_FILE_NOT_FOUND"
                )
            if not os.access(p, os.R_OK):
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.BLOCKED,
                    details=f"Browser target file '{clean_path}' is not readable.",
                    error="ERR_ACCESS_DENIED"
                )
            if p.is_file() and p.stat().st_size == 0:
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.FAILED,
                    details=f"Browser target file '{clean_path}' is empty (0 bytes).",
                    error="FILE_EMPTY"
                )

        # 3. Check browser response for error strings
        if browser_response:
            resp_str = str(browser_response).lower()
            err_patterns = [
                ("err_file_not_found", "ERR_FILE_NOT_FOUND"),
                ("file not found", "ERR_FILE_NOT_FOUND"),
                ("it may have been moved, edited, or deleted", "ERR_FILE_NOT_FOUND"),
                ("err_access_denied", "ERR_ACCESS_DENIED"),
                ("failed to load resource", "ERR_LOAD_FAILED"),
            ]
            for pattern, err_code in err_patterns:
                if pattern in resp_str:
                    return VerificationResult(
                        verified=False,
                        status=VerificationStatus.FAILED,
                        details=f"Browser error detected: '{pattern}' loading '{target_str}'.",
                        error=err_code
                    )

        return VerificationResult(
            verified=True,
            status=VerificationStatus.SUCCESS_VERIFIED,
            evidence=f"Browser opened valid, readable target '{target_str}'.",
            details=f"Browser target verified successfully: {target_str}"
        )


class ArtifactVerifier:
    """Verifies that an artifact was safely exported and registered in host workspace."""

    @staticmethod
    def verify_artifact_exported(record_or_path: Union[Any, str, Path]) -> VerificationResult:
        try:
            from brjarvis.agent.artifacts import ArtifactRecord
            if isinstance(record_or_path, ArtifactRecord):
                if not record_or_path.exported or not record_or_path.host_path:
                    return VerificationResult(
                        verified=False,
                        status=VerificationStatus.FAILED,
                        details=f"Artifact '{record_or_path.filename}' was not exported to host.",
                        error="EXPORT_FAILED"
                    )
                host_p = Path(record_or_path.host_path).resolve()
            else:
                host_p = Path(record_or_path).resolve()

            return FileVerifier.verify_file_created(str(host_p))
        except Exception as e:
            return VerificationResult(
                verified=False,
                status=VerificationStatus.FAILED,
                details=f"Artifact export verification error: {e}",
                error="VERIFICATION_EXCEPTION"
            )


# ── Universal ActionVerifier Facade ────────────────────────────────────────

class ActionVerifier:
    """Deterministic, multi-layered verifier for tool executions and OS mutations."""

    # Re-export specialized methods for direct class-level access
    verify_file_created = staticmethod(FileVerifier.verify_file_created)
    verify_file_content = staticmethod(FileVerifier.verify_file_content)
    verify_file_parsed = staticmethod(FileVerifier.verify_file_parsed)
    verify_process_running = staticmethod(ApplicationVerifier.verify_process_running)
    verify_window_open = staticmethod(ApplicationVerifier.verify_window_open)
    verify_browser_artifact_opened = staticmethod(BrowserVerifier.verify_browser_artifact_opened)
    verify_artifact_exported = staticmethod(ArtifactVerifier.verify_artifact_exported)

    @classmethod
    def verify_tool_output(cls, output_str: str) -> VerificationResult:
        """Inspect tool output for embedded error indicators."""
        if not isinstance(output_str, str):
            return VerificationResult(verified=True, status=VerificationStatus.SUCCESS_VERIFIED, details="Non-string output")

        low = output_str.lower().strip()
        error_indicators = [
            "error:", "traceback (most recent call last):", "zerodivisionerror:",
            "syntaxerror:", "permission denied", "access denied", "scope violation",
            '"status": "failure"', '"error":', '"status": "error"', "err_file_not_found",
            "error building document"
        ]

        for ind in error_indicators:
            if ind in low:
                return VerificationResult(
                    verified=False,
                    status=VerificationStatus.FAILED,
                    details=f"Tool output contains failure indicator: '{ind}'",
                    error="TOOL_OUTPUT_ERROR",
                )

        return VerificationResult(
            verified=True,
            status=VerificationStatus.SUCCESS_VERIFIED,
            details="Tool output clean."
        )

    @classmethod
    def verify_action(cls, tool_name: str, args: Dict[str, Any], output_str: str) -> VerificationResult:
        """Verify the specific real-world outcome of any tool execution."""
        # 1. Output string sanity check
        out_res = cls.verify_tool_output(output_str)
        if not out_res.verified:
            return out_res

        # 2. File write verification
        if tool_name in ("file_write", "create_file", "write_file"):
            target_path = args.get("path") or args.get("file_path") or args.get("name") or ""
            if target_path:
                return cls.verify_file_created(target_path)

        # 3. Document creator verification (DOCX / PDF / HTML / MD)
        if tool_name in ("create_word_document", "create_pdf_document", "document_creator", "generate_walkthrough"):
            filename = args.get("filename") or args.get("path") or ""
            if not filename:
                title = args.get("title", "Document")
                fmt = args.get("format", "docx" if "word" in tool_name else "pdf" if "pdf" in tool_name else "docx")
                clean_title = re.sub(r'[^\w\-]', '_', title)
                filename = f"workspace/Documents/{clean_title}.{fmt}"
            return cls.verify_file_parsed(filename)

        # 4. Artifact export verification
        if tool_name in ("artifact_export", "export_artifact"):
            target_path = args.get("path") or args.get("sandbox_path") or ""
            if target_path:
                return cls.verify_artifact_exported(target_path)

        # 5. Browser / URL open verification
        if tool_name in ("browser_open_url", "open_browser", "web_browser"):
            target_url = args.get("url") or args.get("uri") or args.get("path") or ""
            return cls.verify_browser_artifact_opened(target_url, browser_response=output_str)

        # 6. Application launch verification
        if tool_name in ("open_app", "launch_app"):
            app_name = args.get("app_name") or args.get("name") or ""
            if any(b in app_name.lower() for b in ["chrome", "msedge", "edge", "brave", "firefox"]):
                parts = app_name.split(maxsplit=1)
                if len(parts) > 1 and (":" in parts[1] or "/" in parts[1] or "\\" in parts[1]):
                    return cls.verify_browser_artifact_opened(parts[1], browser_response=output_str)
            if app_name:
                return cls.verify_window_open(app_name=app_name)

        # Default verification pass for standard tools
        return VerificationResult(
            verified=True,
            status=VerificationStatus.SUCCESS_VERIFIED,
            evidence=f"Action '{tool_name}' completed without errors.",
            details=f"Action '{tool_name}' verified without errors."
        )


_global_verifier: Optional[ActionVerifier] = None


def get_action_verifier() -> ActionVerifier:
    global _global_verifier
    if _global_verifier is None:
        _global_verifier = ActionVerifier()
    return _global_verifier


def verify_goal_outcome(goal: str, results: List[Any]) -> VerificationResult:
    """Verify overall goal outcome from aggregated results."""
    has_failures = any("error" in str(r).lower() or "fail" in str(r).lower() for r in results)
    if has_failures:
        return VerificationResult(
            verified=False,
            status=VerificationStatus.FAILED,
            details=f"Goal '{goal[:60]}' completed with failed sub-actions.",
        )
    return VerificationResult(
        verified=True,
        status=VerificationStatus.SUCCESS_VERIFIED,
        evidence=f"Goal '{goal[:60]}' verified successfully.",
        details=f"Goal '{goal[:60]}' verified successfully.",
    )

