# src/brjarvis/tools/doc_tools.py — BR JARVIS Executive Document Generator Engine v3
"""
Automated Executive Document Creator for Microsoft Word (.docx), PDF (.pdf), HTML (.html), and Markdown (.md).
Features Publication-Grade Typography, Cover Pages, Styled Tables, Multi-Line Callout Boxes,
Code Syntax Blocks, Headers, Footers, Non-destructive Path Resolution, Artifact Registration, and Auto-Launching.
"""
from __future__ import annotations

import html
import logging
import os
import re
import shutil
import subprocess
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from .registry import register_tool
from brjarvis.core.paths import paths, get_workspace_manager

logger = logging.getLogger("JARVIS.DocTools")

try:
    import docx  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False


# ── Path Resolution Engine ──────────────────────────────────────────────────
def _resolve_doc_path(filename: str, title: str, fmt: str) -> Path:
    """
    Deterministically resolve target document path without nested workspace/ duplication.

    Rules:
      1. If filename is empty: defaults to paths.DOCUMENTS_DIR / {clean_title}.{fmt}
      2. If filename is absolute: returns resolved Path(filename)
      3. If filename starts with 'workspace/' or 'workspace\\': normalizes under WORKSPACE_ROOT
      4. If filename contains '/' or '\\' (e.g. 'Reports/doc.docx'): resolves under WORKSPACE_ROOT
      5. If filename is a bare name (e.g. 'doc.docx'): resolves under paths.DOCUMENTS_DIR
      6. Ensures proper file extension
    """
    clean_fmt = fmt.lower().lstrip(".").strip() or "docx"
    clean_title = re.sub(r'[^\w\-]', '_', title.strip() or "Document")

    if not filename or not filename.strip():
        target = paths.DOCUMENTS_DIR / f"{clean_title}.{clean_fmt}"
        target.parent.mkdir(parents=True, exist_ok=True)
        return target.resolve()

    raw = filename.strip().replace("\\", "/")

    # Absolute path check
    candidate = Path(raw)
    if candidate.is_absolute():
        if not candidate.suffix:
            candidate = candidate.with_suffix(f".{clean_fmt}")
        candidate.parent.mkdir(parents=True, exist_ok=True)
        return candidate.resolve()

    # Strip redundant leading relative notation
    if raw.startswith("./"):
        raw = raw[2:]

    # Strip leading workspace/ prefix to avoid workspace/workspace/ nesting
    ws_name = paths.WORKSPACE_ROOT.name.lower()
    if raw.lower().startswith(f"{ws_name}/"):
        raw = raw[len(ws_name) + 1:]

    # If subdirectories specified, resolve under WORKSPACE_ROOT; otherwise under DOCUMENTS_DIR
    if "/" in raw:
        target = paths.WORKSPACE_ROOT / raw
    else:
        target = paths.DOCUMENTS_DIR / raw

    if not target.suffix:
        target = target.with_suffix(f".{clean_fmt}")

    target.parent.mkdir(parents=True, exist_ok=True)
    return target.resolve()


def _open_document(saved_path: Path) -> str:
    """Launch the generated document in the host system's default application."""
    target_str = str(saved_path.resolve())
    if sys.platform == "win32":
        try:
            if hasattr(os, "startfile"):
                os.startfile(target_str)
                return f"Auto-launched via Windows Shell: {saved_path.name}"
            subprocess.Popen(["cmd.exe", "/c", "start", "", target_str], shell=False)
            return f"Auto-launched via cmd start: {saved_path.name}"
        except Exception as e:
            return f"Launch notice (win32): {e}"
    elif sys.platform == "darwin":
        try:
            subprocess.Popen(["open", target_str])
            return f"Auto-launched via macOS open: {saved_path.name}"
        except Exception as e:
            return f"Launch notice (darwin): {e}"
    else:
        try:
            subprocess.Popen(["xdg-open", target_str])
            return f"Auto-launched via xdg-open: {saved_path.name}"
        except Exception as e:
            return f"Launch notice (linux): {e}"


# ── XML Styling Helpers for DOCX ──────────────────────────────────────────
def set_cell_background(cell: Any, fill_hex: str):
    """Set the background color of a table cell in hex format (e.g. '1B365D')."""
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_left_border(cell: Any, border_hex: str = '1B365D', border_size_pt: float = 4.0):
    """Set a thick left border and remove top, bottom, and right borders for callout boxes."""
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(int(border_size_pt * 8)))
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_hex)
    tcBorders.append(left)

    for side in ('top', 'bottom', 'right'):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)

    tcPr.append(tcBorders)


def set_cell_margins(cell: Any, top_pt: int = 6, bottom_pt: int = 6, left_pt: int = 10, right_pt: int = 10):
    """Set inner padding margins for table cells."""
    if not _DOCX_AVAILABLE:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top_pt * 14), ('bottom', bottom_pt * 14), ('left', left_pt * 14), ('right', right_pt * 14)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_docx_callout(
    doc: Any,
    text: str,
    callout_type: str = "NOTE",
    border_hex: str = "0284C7",
    fill_hex: str = "F0F9FF"
):
    """Add a styled executive callout box with a colored left border and subtle tinted background."""
    palette = {
        "NOTE": ("0284C7", "F0F9FF", "📌 NOTE"),
        "TIP": ("059669", "F0FDF4", "💡 TIP"),
        "IMPORTANT": ("7C3AED", "FAF5FF", "⭐ IMPORTANT"),
        "WARNING": ("D97706", "FFFBEB", "⚠️ WARNING"),
        "CAUTION": ("DC2626", "FEF2F2", "🛑 CAUTION"),
    }
    b_hex, f_hex, badge = palette.get(callout_type.upper(), (border_hex, fill_hex, "📌 NOTE"))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.5)

    set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=14, right_pt=14)
    set_cell_background(cell, f_hex)
    set_cell_left_border(cell, b_hex, 4.0)

    # Badge paragraph
    p_badge = cell.paragraphs[0]
    p_badge.paragraph_format.space_before = Pt(2)
    p_badge.paragraph_format.space_after = Pt(2)
    r_badge = p_badge.add_run(badge)
    r_badge.font.name = 'Calibri'
    r_badge.font.size = Pt(9.5)
    r_badge.font.bold = True
    r_badge.font.color.rgb = RGBColor(int(b_hex[0:2], 16), int(b_hex[2:4], 16), int(b_hex[4:6], 16))

    # Content paragraphs
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines:
        p_text = cell.add_paragraph()
        p_text.paragraph_format.space_before = Pt(1)
        p_text.paragraph_format.space_after = Pt(2)
        p_text.paragraph_format.line_spacing = 1.15
        _add_paragraph_runs(p_text, line, font_size_pt=10.0, font_color_rgb=(0x33, 0x41, 0x55))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def inject_document_aliases(doc: Any):
    """Inject robust method aliases onto docx Document instance to prevent AI syntax errors."""
    doc.addpagebreak = doc.add_page_break
    doc.add_pagebreak = doc.add_page_break
    doc.pagebreak = doc.add_page_break
    doc.addcallout = lambda text, **kw: add_docx_callout(doc, text, **kw)
    doc.add_callout = lambda text, **kw: add_docx_callout(doc, text, **kw)


# ── Inline Markdown Lexer & Run Generator ──────────────────────────────────
_INLINE_RE = re.compile(
    r'(\*\*[^\*\n]+\*\*|__[^\_\n]+__|(?<!\*)\*[^\*\n]+\*(?!\*)|(?<!_)_[^_\n]+_(?!_)|`[^`\n]+`|\[[^\]\n]+\]\([^\)\n]+\))'
)

def _parse_inline_tokens(text: str) -> List[Tuple[str, str, Optional[str]]]:
    """Tokenize inline markdown into (kind, content, extra) tuples."""
    tokens: List[Tuple[str, str, Optional[str]]] = []
    last_idx = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > last_idx:
            tokens.append(("normal", text[last_idx:match.start()], None))
        raw = match.group(0)
        if (raw.startswith("**") and raw.endswith("**")) or (raw.startswith("__") and raw.endswith("__")):
            tokens.append(("bold", raw[2:-2], None))
        elif raw.startswith("`") and raw.endswith("`"):
            tokens.append(("code", raw[1:-1], None))
        elif raw.startswith("[") and "](" in raw and raw.endswith(")"):
            inner = raw[1:-1]
            txt, link_url = inner.split("](", 1)
            tokens.append(("link", txt, link_url))
        elif (raw.startswith("*") and raw.endswith("*")) or (raw.startswith("_") and raw.endswith("_")):
            tokens.append(("italic", raw[1:-1], None))
        else:
            tokens.append(("normal", raw, None))
        last_idx = match.end()
    if last_idx < len(text):
        tokens.append(("normal", text[last_idx:], None))
    return tokens


def _add_paragraph_runs(
    p: Any,
    text: str,
    font_name: str = 'Calibri',
    font_size_pt: float = 11.0,
    font_color_rgb: Tuple[int, int, int] = (0x22, 0x22, 0x22)
):
    """Parse inline markdown tokens and append styled runs to a DOCX paragraph."""
    tokens = _parse_inline_tokens(text)
    for kind, content, extra in tokens:
        if not content:
            continue
        run = p.add_run(content)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)

        if kind == "bold":
            run.bold = True
            run.font.color.rgb = RGBColor(*font_color_rgb)
        elif kind == "italic":
            run.italic = True
            run.font.color.rgb = RGBColor(*font_color_rgb)
        elif kind == "code":
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size_pt * 0.92)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif kind == "link":
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            run.underline = True
        else:
            run.font.color.rgb = RGBColor(*font_color_rgb)


# ── Advanced DOCX Document Builder ─────────────────────────────────────────
def _build_executive_docx(
    title: str,
    subtitle: str,
    author: str,
    content: str,
    out_path: Path,
    cover_page: bool = True
) -> str:
    """Build a publication-grade Word Document with Cover Page, Callouts, Tables, and Code Blocks."""
    doc = docx.Document()
    inject_document_aliases(doc)

    # Configure Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- COVER PAGE ---
    if cover_page:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(90)
        title_p.paragraph_format.space_after = Pt(12)

        t_run = title_p.add_run(title.upper())
        t_run.font.name = 'Calibri'
        t_run.font.size = Pt(30)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

        if subtitle:
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_p.paragraph_format.space_after = Pt(110)
            s_run = sub_p.add_run(subtitle)
            s_run.font.name = 'Calibri'
            s_run.font.size = Pt(13.5)
            s_run.font.italic = True
            s_run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_p.paragraph_format.space_after = Pt(4)
        a_run = auth_p.add_run(f"Authored by: {author or 'BR JARVIS Autonomous Intelligence'}")
        a_run.font.name = 'Calibri'
        a_run.font.size = Pt(11)
        a_run.font.bold = True
        a_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.paragraph_format.space_after = Pt(20)
        d_run = date_p.add_run(time.strftime("%B %d, %Y • Executive Edition"))
        d_run.font.name = 'Calibri'
        d_run.font.size = Pt(10)
        d_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        doc.add_page_break()
    else:
        h1 = doc.add_heading(title, level=0)
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if subtitle:
            p_sub = doc.add_paragraph()
            r_sub = p_sub.add_run(subtitle)
            r_sub.font.italic = True
            r_sub.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    # --- BODY CONTENT PARSER ---
    lines = content.splitlines()
    i = 0
    in_code_block = False
    code_lines: List[str] = []

    while i < len(lines):
        line = lines[i]
        line_s = line.strip()

        # Handle Code Blocks
        if line_s.startswith("```"):
            if in_code_block:
                # Flush code block
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                c = tbl.cell(0, 0)
                c.width = Inches(6.5)
                set_cell_background(c, '0F172A')  # Slate Dark
                set_cell_margins(c, top_pt=8, bottom_pt=8, left_pt=12, right_pt=12)
                p_code = c.paragraphs[0]
                p_code.paragraph_format.space_before = Pt(2)
                p_code.paragraph_format.space_after = Pt(2)
                p_code.paragraph_format.line_spacing = 1.0
                r_code = p_code.add_run("\n".join(code_lines))
                r_code.font.name = 'Consolas'
                r_code.font.size = Pt(9.5)
                r_code.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)

                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not line_s:
            i += 1
            continue

        # Handle Multi-line Callout Boxes (> [!NOTE] ...)
        if line_s.startswith(">"):
            callout_lines: List[str] = []
            callout_type = "NOTE"
            while i < len(lines) and lines[i].strip().startswith(">"):
                raw_c = lines[i].strip().lstrip("> ").strip()
                if raw_c.startswith("[!") and "]" in raw_c:
                    ctype = raw_c[2:raw_c.find("]")].strip().upper()
                    if ctype in ("NOTE", "TIP", "IMPORTANT", "WARNING", "CAUTION"):
                        callout_type = ctype
                    remainder = raw_c[raw_c.find("]") + 1:].strip()
                    if remainder:
                        callout_lines.append(remainder)
                else:
                    if raw_c:
                        callout_lines.append(raw_c)
                i += 1

            if callout_lines:
                add_docx_callout(doc, "\n".join(callout_lines), callout_type=callout_type)
            continue

        # Handle Markdown Tables
        if "|" in line_s and (line_s.startswith("|") or line_s.endswith("|")):
            table_lines: List[str] = []
            while i < len(lines) and ("|" in lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            parsed_rows: List[List[str]] = []
            for tl in table_lines:
                raw_cells = tl.split("|")
                # Strip leading/trailing empty split fragments if present
                if tl.startswith("|"):
                    raw_cells = raw_cells[1:]
                if tl.endswith("|") and raw_cells:
                    raw_cells = raw_cells[:-1]
                cells = [c.strip() for c in raw_cells]
                # Check if it is a markdown separator row (e.g. |---|---|)
                is_sep = all(all(ch in "-: " for ch in cell) for cell in cells) if cells else False
                if not is_sep and any(c for c in cells):
                    parsed_rows.append(cells)

            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                tbl = doc.add_table(rows=0, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r_idx, row_cells in enumerate(parsed_rows):
                    row = tbl.add_row()
                    for c_idx in range(num_cols):
                        cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell = row.cells[c_idx]
                        set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
                        p_cell = cell.paragraphs[0]
                        p_cell.paragraph_format.space_before = Pt(2)
                        p_cell.paragraph_format.space_after = Pt(2)

                        if r_idx == 0:
                            set_cell_background(cell, '1B365D')  # Navy Header
                            _add_paragraph_runs(p_cell, cell_text, font_color_rgb=(0xFF, 0xFF, 0xFF))
                            for run in p_cell.runs:
                                run.bold = True
                        else:
                            bg_fill = 'F8FAFC' if r_idx % 2 == 1 else 'FFFFFF'
                            set_cell_background(cell, bg_fill)
                            _add_paragraph_runs(p_cell, cell_text, font_color_rgb=(0x22, 0x22, 0x22))

                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # Headings
        if line_s.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(22)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[2:])
            r.font.name = 'Calibri'
            r.font.size = Pt(18)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        elif line_s.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(15)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[3:])
            r.font.name = 'Calibri'
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)  # Teal Accent
        elif line_s.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(11)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[4:])
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif line_s.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line_s[5:])
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Horizontal Divider (--- or ***)
        elif re.match(r'^(?:---|\*\*\*|___)\s*$', line_s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run("─" * 45)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

        # Bullet List items
        elif line_s.startswith("- ") or line_s.startswith("* ") or line_s.startswith("+ "):
            indent_level = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            if indent_level >= 4:
                p.paragraph_format.left_indent = Inches(0.5)
            _add_paragraph_runs(p, line_s[2:])

        # Numbered List items
        elif re.match(r'^\d+\.\s+', line_s):
            match = re.match(r'^(\d+\.\s+)(.*)', line_s)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            run_num = p.add_run(match.group(1))
            run_num.font.name = 'Calibri'
            run_num.font.bold = True
            run_num.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            _add_paragraph_runs(p, match.group(2))

        # Regular Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.line_spacing = 1.15
            _add_paragraph_runs(p, line_s)

        i += 1

    # Flush unclosed code block if necessary
    if in_code_block and code_lines:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.cell(0, 0)
        c.width = Inches(6.5)
        set_cell_background(c, '0F172A')
        p_code = c.paragraphs[0]
        r_code = p_code.add_run("\n".join(code_lines))
        r_code.font.name = 'Consolas'
        r_code.font.size = Pt(9.5)
        r_code.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)


def _sanitize_pdf_text(text: str) -> str:
    """Normalize unicode characters for PDF core fonts (Helvetica / Courier / Times)."""
    if not text:
        return ""
    replacements = {
        "\u2022": "*",   # Bullet •
        "\u2023": "*",   # Triangular bullet
        "\u2043": "*",   # Hyphen bullet
        "\u2013": "-",   # En dash –
        "\u2014": "--",  # Em dash —
        "\u2015": "--",  # Horizontal bar
        "\u2018": "'",   # Left single quote ‘
        "\u2019": "'",   # Right single quote ’
        "\u201a": "'",   # Single low-9 quote
        "\u201c": '"',   # Left double quote “
        "\u201d": '"',   # Right double quote ”
        "\u201e": '"',   # Double low-9 quote
        "\u2026": "...", # Ellipsis …
        "\u00a0": " ",   # Non-breaking space
        "\u2192": "->",  # Right arrow
        "\u2190": "<-",  # Left arrow
        "\u2713": "[v]", # Checkmark
        "\u2714": "[v]", # Heavy checkmark
        "\u2717": "[x]", # Cross mark
        "\u2718": "[x]", # Heavy cross mark
        "\u26a0": "[!]", # Warning sign
        "\u2728": "*",   # Sparkles
        "\u26a1": "[*]", # Lightning ⚡
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    return text.encode("latin-1", "replace").decode("latin-1")


# ── Advanced PDF Builder (FPDF2) ───────────────────────────────────────────
if _FPDF_AVAILABLE:
    class ExecutivePDF(FPDF):
        """Custom PDF class with running header, footer, and page numbering."""

        def __init__(self, doc_title: str = "Executive Document"):
            super().__init__()
            self.doc_title = _sanitize_pdf_text(doc_title)

        def header(self):
            if self.page_no() > 1:
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(140, 150, 160)
                self.cell(0, 6, self.doc_title, align="R")
                self.ln(8)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(140, 150, 160)
            self.cell(0, 10, f"Page {self.page_no()} | BR JARVIS Autonomous Executive Intelligence", align="C")


def _build_executive_pdf(
    title: str,
    subtitle: str,
    author: str,
    content: str,
    out_path: Path
) -> str:
    """Build a styled publication PDF document using FPDF2 with Tables, Callouts, and Code Blocks."""
    pdf = ExecutivePDF(doc_title=title)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # --- Title Banner ---
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(27, 54, 93)  # Navy
    clean_title = _sanitize_pdf_text(title)
    pdf.multi_cell(pdf.epw, 9, clean_title, align="L", new_x="LMARGIN", new_y="NEXT")

    if subtitle:
        pdf.set_font("Helvetica", "I", 11)
        pdf.set_text_color(100, 116, 139)  # Slate
        clean_sub = _sanitize_pdf_text(subtitle)
        pdf.multi_cell(pdf.epw, 6, clean_sub, align="L", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(71, 85, 105)
    meta_line = _sanitize_pdf_text(f"Author: {author or 'BR JARVIS'}  |  Date: {time.strftime('%B %d, %Y')}")
    pdf.cell(pdf.epw, 6, meta_line, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # Thin Divider Rule
    pdf.set_draw_color(203, 213, 225)
    pdf.set_line_width(0.4)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + pdf.epw, pdf.get_y())
    pdf.ln(5)

    # --- Body Parser ---
    lines = content.splitlines()
    i = 0
    in_code = False
    code_buf: List[str] = []

    while i < len(lines):
        line = lines[i]
        line_s = line.strip()

        # Handle Code Block
        if line_s.startswith("```"):
            if in_code:
                # Flush code block
                pdf.set_font("Courier", size=8.5)
                pdf.set_fill_color(241, 245, 249)  # Slate light
                pdf.set_draw_color(203, 213, 225)
                pdf.set_text_color(30, 41, 59)
                code_text = _sanitize_pdf_text("\n".join(code_buf))
                pdf.multi_cell(pdf.epw, 4.5, code_text, fill=True, border=1, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line_s:
            pdf.ln(2)
            i += 1
            continue

        # Handle Callout Box (> [!NOTE] ...)
        if line_s.startswith(">"):
            c_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                raw_c = lines[i].strip().lstrip("> ").strip()
                if raw_c.startswith("[!") and "]" in raw_c:
                    raw_c = raw_c[raw_c.find("]") + 1:].strip()
                if raw_c:
                    c_lines.append(raw_c)
                i += 1

            if c_lines:
                pdf.set_font("Helvetica", "I", 9.5)
                pdf.set_fill_color(240, 249, 255)  # Sky tint
                pdf.set_draw_color(2, 132, 199)   # Sky border
                pdf.set_text_color(15, 23, 42)
                c_text = _sanitize_pdf_text(" * " + "\n * ".join(c_lines))
                pdf.multi_cell(pdf.epw, 5, c_text, fill=True, border=1, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
            continue

        # Handle Markdown Table
        if "|" in line_s and (line_s.startswith("|") or line_s.endswith("|")):
            t_lines: List[str] = []
            while i < len(lines) and ("|" in lines[i].strip()):
                t_lines.append(lines[i].strip())
                i += 1

            t_rows: List[List[str]] = []
            for tl in t_lines:
                raw_cells = tl.split("|")
                if tl.startswith("|"):
                    raw_cells = raw_cells[1:]
                if tl.endswith("|") and raw_cells:
                    raw_cells = raw_cells[:-1]
                cells = [_sanitize_pdf_text(c.strip()) for c in raw_cells]
                is_sep = all(all(ch in "-: " for ch in cell) for cell in cells) if cells else False
                if not is_sep and any(c for c in cells):
                    t_rows.append(cells)

            if t_rows and hasattr(pdf, "table"):
                try:
                    with pdf.table(line_height=5.5) as table:
                        for r_idx, row_cells in enumerate(t_rows):
                            row = table.row()
                            for c_txt in row_cells:
                                row.cell(c_txt)
                    pdf.ln(3)
                except Exception as tbl_err:
                    logger.debug("PDF table fallback: %s", tbl_err)
                    pdf.set_font("Helvetica", size=9)
                    for r in t_rows:
                        pdf.cell(pdf.epw, 5, " | ".join(r), new_x="LMARGIN", new_y="NEXT")
            continue

        # Headings
        if line_s.startswith("# "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(27, 54, 93)
            pdf.multi_cell(pdf.epw, 7, _sanitize_pdf_text(line_s[2:]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif line_s.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(13, 148, 136)
            pdf.multi_cell(pdf.epw, 6, _sanitize_pdf_text(line_s[3:]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif line_s.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(51, 65, 85)
            pdf.multi_cell(pdf.epw, 5, _sanitize_pdf_text(line_s[4:]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif line_s.startswith("- ") or line_s.startswith("* "):
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(30, 41, 59)
            bullet_text = _sanitize_pdf_text("  *  " + line_s[2:])
            pdf.multi_cell(pdf.epw, 5, bullet_text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif re.match(r'^\d+\.\s+', line_s):
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(pdf.epw, 5, _sanitize_pdf_text("   " + line_s), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        else:
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(30, 41, 59)
            clean_p = _sanitize_pdf_text(line_s)
            pdf.multi_cell(pdf.epw, 5, clean_p, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))
    return str(out_path)


# ── Advanced Glassmorphism HTML Builder ────────────────────────────────────
def _format_inline_html(text: str) -> str:
    """Escape and format inline markdown formatting for HTML."""
    tokens = _parse_inline_tokens(text)
    out_parts: List[str] = []
    for kind, content, extra in tokens:
        escaped = html.escape(content)
        if kind == "bold":
            out_parts.append(f"<strong>{escaped}</strong>")
        elif kind == "italic":
            out_parts.append(f"<em>{escaped}</em>")
        elif kind == "code":
            out_parts.append(f"<code>{escaped}</code>")
        elif kind == "link":
            url_esc = html.escape(extra or "#")
            out_parts.append(f"<a href='{url_esc}' target='_blank' rel='noopener'>{escaped}</a>")
        else:
            out_parts.append(escaped)
    return "".join(out_parts)


def _build_executive_html(
    title: str,
    subtitle: str,
    author: str,
    content: str,
    out_path: Path
) -> str:
    """Build a modern responsive Glassmorphism HTML document with print styling."""
    title_esc = html.escape(title)
    sub_esc = html.escape(subtitle or "Executive Intelligence Document")
    auth_esc = html.escape(author or "BR JARVIS Autonomous Core")

    html_lines = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "<meta charset='UTF-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        f"<title>{title_esc}</title>",
        "<link rel='preconnect' href='https://fonts.googleapis.com'>",
        "<link rel='preconnect' href='https://fonts.gstatic.com' crossorigin>",
        "<link href='https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap' rel='stylesheet'>",
        "<style>",
        "  :root {",
        "    --bg-dark: #090d16;",
        "    --card-bg: rgba(23, 32, 51, 0.85);",
        "    --card-border: rgba(255, 255, 255, 0.08);",
        "    --accent-cyan: #38bdf8;",
        "    --accent-indigo: #818cf8;",
        "    --accent-teal: #2dd4bf;",
        "    --text-primary: #f8fafc;",
        "    --text-secondary: #94a3b8;",
        "    --text-body: #cbd5e1;",
        "  }",
        "  * { box-sizing: border-box; }",
        "  body { font-family: 'Inter', system-ui, -apple-system, sans-serif; background: var(--bg-dark); background-image: radial-gradient(at 0% 0%, rgba(56, 189, 248, 0.08) 0px, transparent 50%), radial-gradient(at 100% 100%, rgba(129, 140, 248, 0.08) 0px, transparent 50%); color: var(--text-primary); margin: 0; padding: 48px 24px; min-height: 100vh; line-height: 1.65; }",
        "  .container { max-width: 960px; margin: 0 auto; background: var(--card-bg); backdrop-filter: blur(20px); border: 1px solid var(--card-border); border-radius: 20px; padding: 56px; box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.6); }",
        "  .header-badge { display: inline-flex; align-items: center; gap: 8px; background: rgba(56, 189, 248, 0.12); border: 1px solid rgba(56, 189, 248, 0.25); color: var(--accent-cyan); font-size: 0.8rem; font-weight: 600; padding: 4px 12px; border-radius: 9999px; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 16px; }",
        "  h1 { font-size: 2.5rem; font-weight: 800; line-height: 1.2; margin: 0 0 12px 0; background: linear-gradient(135deg, #ffffff 0%, #cbd5e1 50%, var(--accent-cyan) 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }",
        "  .subtitle-bar { display: flex; flex-wrap: wrap; justify-content: space-between; align-items: center; border-bottom: 1px solid rgba(255, 255, 255, 0.08); padding-bottom: 20px; margin-bottom: 36px; color: var(--text-secondary); font-size: 0.95rem; }",
        "  h2 { font-size: 1.5rem; font-weight: 700; color: var(--accent-cyan); margin-top: 36px; margin-bottom: 12px; border-bottom: 1px solid rgba(56, 189, 248, 0.15); padding-bottom: 8px; }",
        "  h3 { font-size: 1.25rem; font-weight: 600; color: var(--accent-indigo); margin-top: 24px; margin-bottom: 8px; }",
        "  h4 { font-size: 1.05rem; font-weight: 600; color: #e2e8f0; margin-top: 18px; margin-bottom: 6px; }",
        "  p { color: var(--text-body); font-size: 1.02rem; margin: 12px 0; }",
        "  ul, ol { color: var(--text-body); font-size: 1.02rem; padding-left: 24px; margin: 12px 0; }",
        "  li { margin-bottom: 6px; }",
        "  .callout { background: rgba(56, 189, 248, 0.08); border-left: 4px solid var(--accent-cyan); padding: 18px 24px; border-radius: 6px 12px 12px 6px; margin: 24px 0; }",
        "  .callout-badge { font-weight: 700; font-size: 0.8rem; letter-spacing: 0.05em; text-transform: uppercase; color: var(--accent-cyan); margin-bottom: 4px; }",
        "  .table-wrapper { width: 100%; overflow-x: auto; margin: 24px 0; border-radius: 12px; border: 1px solid rgba(255, 255, 255, 0.08); }",
        "  table { width: 100%; border-collapse: collapse; text-align: left; font-size: 0.95rem; }",
        "  th { background: #1e293b; color: #f1f5f9; padding: 14px 18px; font-weight: 600; border-bottom: 2px solid #334155; }",
        "  td { padding: 12px 18px; border-bottom: 1px solid rgba(255, 255, 255, 0.05); color: #e2e8f0; }",
        "  tr:nth-child(even) { background: rgba(255, 255, 255, 0.02); }",
        "  tr:hover { background: rgba(56, 189, 248, 0.04); }",
        "  pre { background: #050811; border: 1px solid #1e293b; border-radius: 10px; padding: 20px; overflow-x: auto; font-family: 'JetBrains Mono', Consolas, monospace; font-size: 0.9rem; color: #e2e8f0; margin: 20px 0; }",
        "  code { font-family: 'JetBrains Mono', Consolas, monospace; background: rgba(255, 255, 255, 0.08); padding: 2px 6px; border-radius: 4px; font-size: 0.9em; color: #38bdf8; }",
        "  pre code { background: transparent; padding: 0; color: inherit; }",
        "  a { color: var(--accent-cyan); text-decoration: none; border-bottom: 1px dotted rgba(56, 189, 248, 0.5); }",
        "  a:hover { border-bottom-style: solid; }",
        "  @media print { body { background: white; color: black; padding: 0; } .container { box-shadow: none; border: none; padding: 0; background: white; } }",
        "</style>",
        "</head>",
        "<body>",
        "<div class='container'>",
        "  <div class='header-badge'>⚡ BR JARVIS Executive Report</div>",
        f"  <h1>{title_esc}</h1>",
        f"  <div class='subtitle-bar'><span>{sub_esc}</span><span>{auth_esc} • {time.strftime('%b %d, %Y')}</span></div>",
    ]

    lines = content.splitlines()
    i = 0
    in_code = False
    code_buf: List[str] = []
    in_ul = False
    in_ol = False

    def _close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            html_lines.append("</ul>")
            in_ul = False
        if in_ol:
            html_lines.append("</ol>")
            in_ol = False

    while i < len(lines):
        line = lines[i]
        line_s = line.strip()

        # Handle Code Block
        if line_s.startswith("```"):
            _close_lists()
            if in_code:
                code_raw = html.escape("\n".join(code_buf))
                html_lines.append(f"<pre><code>{code_raw}</code></pre>")
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line_s:
            _close_lists()
            i += 1
            continue

        # Handle Callout Box (> [!NOTE] ...)
        if line_s.startswith(">"):
            _close_lists()
            callout_lines: List[str] = []
            badge_label = "NOTE"
            while i < len(lines) and lines[i].strip().startswith(">"):
                raw_c = lines[i].strip().lstrip("> ").strip()
                if raw_c.startswith("[!") and "]" in raw_c:
                    badge_label = raw_c[2:raw_c.find("]")].strip().upper()
                    remainder = raw_c[raw_c.find("]") + 1:].strip()
                    if remainder:
                        callout_lines.append(remainder)
                else:
                    if raw_c:
                        callout_lines.append(raw_c)
                i += 1

            c_body = "<br>".join(_format_inline_html(l) for l in callout_lines)
            html_lines.append(
                f"<div class='callout'><div class='callout-badge'>📌 {html.escape(badge_label)}</div><p>{c_body}</p></div>"
            )
            continue

        # Handle Markdown Tables
        if "|" in line_s and (line_s.startswith("|") or line_s.endswith("|")):
            _close_lists()
            t_lines: List[str] = []
            while i < len(lines) and ("|" in lines[i].strip()):
                t_lines.append(lines[i].strip())
                i += 1

            t_rows: List[List[str]] = []
            for tl in t_lines:
                raw_cells = tl.split("|")
                if tl.startswith("|"):
                    raw_cells = raw_cells[1:]
                if tl.endswith("|") and raw_cells:
                    raw_cells = raw_cells[:-1]
                cells = [c.strip() for c in raw_cells]
                is_sep = all(all(ch in "-: " for ch in cell) for cell in cells) if cells else False
                if not is_sep and any(c for c in cells):
                    t_rows.append(cells)

            if t_rows:
                html_lines.append("<div class='table-wrapper'><table>")
                # Header row
                html_lines.append("<thead><tr>")
                for hc in t_rows[0]:
                    html_lines.append(f"<th>{_format_inline_html(hc)}</th>")
                html_lines.append("</tr></thead><tbody>")
                # Data rows
                for dr in t_rows[1:]:
                    html_lines.append("<tr>")
                    for dc in dr:
                        html_lines.append(f"<td>{_format_inline_html(dc)}</td>")
                    html_lines.append("</tr>")
                html_lines.append("</tbody></table></div>")
            continue

        # Headings
        if line_s.startswith("# "):
            _close_lists()
            html_lines.append(f"<h2>{_format_inline_html(line_s[2:])}</h2>")
        elif line_s.startswith("## "):
            _close_lists()
            html_lines.append(f"<h2>{_format_inline_html(line_s[3:])}</h2>")
        elif line_s.startswith("### "):
            _close_lists()
            html_lines.append(f"<h3>{_format_inline_html(line_s[4:])}</h3>")
        elif line_s.startswith("#### "):
            _close_lists()
            html_lines.append(f"<h4>{_format_inline_html(line_s[5:])}</h4>")

        # Bullet lists
        elif line_s.startswith("- ") or line_s.startswith("* ") or line_s.startswith("+ "):
            if not in_ul:
                _close_lists()
                html_lines.append("<ul>")
                in_ul = True
            html_lines.append(f"<li>{_format_inline_html(line_s[2:])}</li>")

        # Numbered lists
        elif re.match(r'^\d+\.\s+', line_s):
            match = re.match(r'^\d+\.\s+(.*)', line_s)
            if not in_ol:
                _close_lists()
                html_lines.append("<ol>")
                in_ol = True
            html_lines.append(f"<li>{_format_inline_html(match.group(1))}</li>")

        # Normal paragraph
        else:
            _close_lists()
            html_lines.append(f"<p>{_format_inline_html(line_s)}</p>")

        i += 1

    _close_lists()
    if in_code and code_buf:
        code_raw = html.escape("\n".join(code_buf))
        html_lines.append(f"<pre><code>{code_raw}</code></pre>")

    html_lines.extend(["</div>", "</body>", "</html>"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(html_lines), encoding="utf-8")
    return str(out_path)


# ── Universal Document Creator Tool ────────────────────────────────────────
@register_tool(
    name="document_creator",
    description="Universal Executive Document Engine. Creates styled Word (.docx), PDF (.pdf), HTML (.html), and Markdown (.md) documents with Cover Pages, Styled Tables, Callouts, Code Blocks, and Auto-Launch.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title of the document"},
            "subtitle": {"type": "string", "description": "Subtitle or tagline for the document"},
            "author": {"type": "string", "description": "Author name (default: BR JARVIS AI)"},
            "content": {"type": "string", "description": "Main text content in structured Markdown (headings, bullets, tables, callouts)"},
            "filename": {"type": "string", "description": "Target filename or relative path (e.g., Reports/Startup_Book.docx or documents/analysis.pdf)"},
            "format": {"type": "string", "description": "Output format: docx | pdf | html | md (default: docx)"},
            "cover_page": {"type": "boolean", "description": "Whether to include an executive cover page (default: true)"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch the generated file (default: true)"}
        },
        "required": ["title", "content"]
    }
)
def document_creator(args: dict) -> str:
    """Universal Executive Document Engine."""
    title = str(args.get("title") or "Document").strip()
    subtitle = str(args.get("subtitle") or "").strip()
    author = str(args.get("author") or "BR JARVIS Autonomous Intelligence").strip()
    content = str(args.get("content") or "").strip()
    fmt = str(args.get("format") or "docx").lower().strip()
    cover_page = bool(args.get("cover_page", True))
    auto_open = bool(args.get("auto_open", True))
    filename = str(args.get("filename") or "").strip()

    out_path = _resolve_doc_path(filename, title, fmt)

    try:
        if fmt == "docx":
            if not _DOCX_AVAILABLE:
                return "Error: 'python-docx' library is not installed."
            saved_path = _build_executive_docx(title, subtitle, author, content, out_path, cover_page=cover_page)
        elif fmt == "pdf":
            if not _FPDF_AVAILABLE:
                return "Error: 'fpdf2' library is not installed."
            saved_path = _build_executive_pdf(title, subtitle, author, content, out_path)
        elif fmt == "html":
            saved_path = _build_executive_html(title, subtitle, author, content, out_path)
        elif fmt in ("md", "markdown"):
            out_path.parent.mkdir(parents=True, exist_ok=True)
            md_text = f"# {title}\n\n*{subtitle}*\n\n**Author:** {author}\n\n---\n\n{content}\n"
            out_path.write_text(md_text, encoding="utf-8")
            saved_path = str(out_path)
        else:
            return f"Error: Unsupported format '{fmt}'. Choose from docx, pdf, html, md."
    except Exception as e:
        logger.exception("Document creation error: %s", e)
        return f"Error building document ({fmt}): {e}"

    # Verify generated document
    from agent.verifier import ActionVerifier
    v_res = ActionVerifier.verify_file_parsed(str(saved_path))
    if not v_res.verified:
        return f"Document created but verification failed: {v_res.details}"

    # Register in ArtifactManager
    try:
        from agent.artifacts import get_artifact_manager
        mgr = get_artifact_manager()
        mgr.export_sandbox_artifact(saved_path, custom_filename=out_path.name)
    except Exception as art_err:
        logger.debug("Artifact registration note: %s", art_err)

    open_status = "auto-open disabled"
    if auto_open:
        open_status = _open_document(Path(saved_path))

    return f"⚡ [SUCCESS_VERIFIED] Created Executive Document ({fmt.upper()}): '{saved_path}' | Evidence: {v_res.evidence} | Viewer: {open_status}"


# ── Backward Compatibility Tool Wrappers ──────────────────────────────────
@register_tool(
    name="create_word_document",
    description="Create a formatted Microsoft Word (.docx) document with cover page, headers, tables, callouts, and auto-launch.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main document text content or markdown"},
            "filename": {"type": "string", "description": "Output filename ending in .docx"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch Word"}
        },
        "required": ["title", "content"]
    }
)
def create_word_document(args: dict) -> str:
    """Create Word document via document_creator engine."""
    args_copy = dict(args) if isinstance(args, dict) else {}
    args_copy["format"] = "docx"
    if not args_copy.get("content"):
        args_copy["content"] = args_copy.get("title", "Document Content")
    return document_creator(args_copy)


@register_tool(
    name="create_pdf_document",
    description="Create a formatted PDF (.pdf) document and auto-launch in default PDF viewer.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main text content or markdown"},
            "filename": {"type": "string", "description": "Output filename ending in .pdf"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch PDF viewer"}
        },
        "required": ["title", "content"]
    }
)
def create_pdf_document(args: dict) -> str:
    """Create PDF document via document_creator engine."""
    args_copy = dict(args) if isinstance(args, dict) else {}
    args_copy["format"] = "pdf"
    if not args_copy.get("content"):
        args_copy["content"] = args_copy.get("title", "Document Content")
    return document_creator(args_copy)


@register_tool(
    name="generate_project_product_analysis",
    description="Generate a complete Product Analysis Report for B.R. JARVIS as Word (.docx) and PDF (.pdf) documents and auto-open them.",
    parameters={"type": "object", "properties": {}}
)
def generate_project_product_analysis(args: dict) -> str:
    """Generate complete Product Analysis report for B.R. JARVIS in Word & PDF formats."""
    doc_title = "Product Analysis Document: B.R. JARVIS"
    doc_text = """
# Product Analysis Document: B.R. JARVIS

## 1. Executive Summary
- **Product Name:** B.R. JARVIS (Advanced Agentic AI Operating System)
- **Identity:** Ultra-fast autonomous AI assistant designed for decisive multi-step reasoning, pair programming, and live desktop visual control.
- **Mission:** To eliminate conversational filler and deliver instant, high-precision software engineering actions.

> [!NOTE]
> B.R. JARVIS operates local-first with zero token latency optimizations and full execution autonomy.

## 2. Product Vision & Value Proposition
B.R. JARVIS shifts the paradigm from static autocomplete AI to an autonomous senior developer:
- **Zero-Filler Directive:** Delivers immediate, high-signal task resolution with minimal output tokens.
- **Local Gateway Integration:** Operates with unlimited request quotas via local gateway http://127.0.0.1:8045/v1.

## 3. Core Features & Capabilities
- **0-Token Intent Engine:** Executes common app launches, browser navigation, and diagnostics in 0ms with zero token cost.
- **Live OS Visual Controller:** Performs real-time desktop visual grounded automation (2160x1440 screen resolution).
- **Multi-Tab Excel & Document Generator:** Automatically creates formatted .xlsx, .docx, and .pdf analytical reports.

## 4. Architecture & Subsystems Inventory
| Subsystem | Components | Function |
| --- | --- | --- |
| Core Kernel | Native C FNV-1a bridge, EventBus runtime | Fast message routing |
| Action Suite | Live OS Control, Computer Control, RAG Library | Desktop & web automation |
| Tool Registry | 140+ registered tools | Plugin architecture |
| UI HUD | Tkinter glassmorphism display | Real-time status feedback |

## 5. Security & Execution Safety
- Local-first workspace execution with absolute path validation.
- AST compilation checks and security vulnerability scanning.
"""
    res_word = document_creator({
        "title": doc_title,
        "content": doc_text,
        "filename": "Reports/JARVIS_Product_Analysis.docx",
        "format": "docx",
        "auto_open": True
    })
    res_pdf = document_creator({
        "title": doc_title,
        "content": doc_text,
        "filename": "Reports/JARVIS_Product_Analysis.pdf",
        "format": "pdf",
        "auto_open": True
    })

    return f"{res_word}\n{res_pdf}"


@register_tool(
    name="generate_walkthrough",
    description="Generate a rich GitHub-flavored Markdown Walkthrough document (walkthrough.md) documenting technical changes, verification results, and file links.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title of the walkthrough"},
            "summary": {"type": "string", "description": "High-level summary of work accomplished"},
            "changes": {"type": "string", "description": "Detailed description or markdown list of changes made"},
            "verification": {"type": "string", "description": "Verification steps and automated test results"},
            "filename": {"type": "string", "description": "Target filename, default is walkthrough.md"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-open the generated walkthrough file"}
        },
        "required": ["title", "changes"]
    }
)
def generate_walkthrough(args: dict) -> str:
    """Generate a GitHub-flavored markdown walkthrough document."""
    title = str(args.get("title") or "Task Walkthrough").strip()
    summary = str(args.get("summary") or "").strip()
    changes = str(args.get("changes") or "").strip()
    verification = str(args.get("verification") or "").strip()
    filename = str(args.get("filename") or "walkthrough.md").strip()
    auto_open = bool(args.get("auto_open", False))

    content_lines = [f"# Walkthrough — {title}\n"]
    if summary:
        content_lines.append(f"{summary}\n")
    content_lines.append("## Changes Made\n")
    content_lines.append(f"{changes}\n")
    if verification:
        content_lines.append("## Verification Results\n")
        content_lines.append(f"{verification}\n")

    full_md = "\n".join(content_lines)
    out_path = _resolve_doc_path(filename, title, "md")

    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(full_md, encoding="utf-8")
    except Exception as e:
        return f"Error writing walkthrough document: {e}"

    if auto_open:
        _open_document(out_path)

    file_uri = out_path.as_uri()
    return f"⚡ Generated Walkthrough document successfully: [{filename}]({file_uri})"
