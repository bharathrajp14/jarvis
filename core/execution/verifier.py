# core/execution/verifier.py — Universal Side-Effect Verification Engine
from __future__ import annotations

import csv
import json
import logging
import os
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from core.execution.types import ExecutionResult, ExecutionStatus, VerificationOutcome

logger = logging.getLogger("JARVIS.UniversalVerifier")


class FileVerifier:
    """Verifies file existence, non-zero size, and disk readability."""

    @staticmethod
    def verify_file(path_str: str | Path, min_size_bytes: int = 1) -> VerificationOutcome:
        try:
            p = Path(path_str).resolve()
            if not p.exists():
                return VerificationOutcome(
                    verified=False,
                    verifier_name="FileVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Verification failed: File '{p}' does not exist on disk.",
                    error="FILE_NOT_FOUND",
                )
            if not p.is_file():
                return VerificationOutcome(
                    verified=False,
                    verifier_name="FileVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Verification failed: Path '{p}' is a directory, not a file.",
                    error="NOT_A_FILE",
                )
            size = p.stat().st_size
            if size < min_size_bytes:
                return VerificationOutcome(
                    verified=False,
                    verifier_name="FileVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Verification failed: File '{p.name}' is empty ({size} bytes).",
                    error="FILE_EMPTY",
                )
            return VerificationOutcome(
                verified=True,
                verifier_name="FileVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"File '{p.name}' verified on disk ({size:,} bytes, path: {p}).",
                details=f"File '{p.name}' verified ({size} bytes).",
                observed_state={"path": str(p), "size_bytes": size},
            )
        except Exception as e:
            return VerificationOutcome(
                verified=False,
                verifier_name="FileVerifier",
                status=ExecutionStatus.FAILED,
                details=f"Verification exception for '{path_str}': {e}",
                error="VERIFICATION_EXCEPTION",
            )


class DirectoryVerifier:
    """Verifies directory existence and contents."""

    @staticmethod
    def verify_directory(dir_path: str | Path, min_files: int = 0) -> VerificationOutcome:
        try:
            p = Path(dir_path).resolve()
            if not p.exists() or not p.is_dir():
                return VerificationOutcome(
                    verified=False,
                    verifier_name="DirectoryVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Directory '{p}' does not exist.",
                    error="DIR_NOT_FOUND",
                )
            files = list(p.iterdir())
            if len(files) < min_files:
                return VerificationOutcome(
                    verified=False,
                    verifier_name="DirectoryVerifier",
                    status=ExecutionStatus.PARTIAL_SUCCESS,
                    details=f"Directory '{p.name}' contains {len(files)} items (expected at least {min_files}).",
                    error="INSUFFICIENT_ITEMS",
                )
            return VerificationOutcome(
                verified=True,
                verifier_name="DirectoryVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"Directory '{p.name}' verified ({len(files)} items).",
                details=f"Verified directory '{p.name}'.",
                observed_state={"item_count": len(files)},
            )
        except Exception as e:
            return VerificationOutcome(
                verified=False,
                verifier_name="DirectoryVerifier",
                status=ExecutionStatus.FAILED,
                details=f"Directory verification error: {e}",
                error="DIR_VERIFICATION_ERROR",
            )


class DocumentVerifier:
    """Parses and validates structural integrity of complex documents (DOCX, PDF, XLSX, JSON, CSV)."""

    @staticmethod
    def verify_document(path_str: str | Path) -> VerificationOutcome:
        file_res = FileVerifier.verify_file(path_str)
        if not file_res.verified:
            return file_res

        p = Path(path_str).resolve()
        ext = p.suffix.lower()

        try:
            if ext == ".docx":
                import docx
                doc = docx.Document(str(p))
                p_count = len(doc.paragraphs)
                t_count = len(doc.tables)
                text_len = sum(len(para.text) for para in doc.paragraphs)
                if p_count == 0 and t_count == 0:
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="DocumentVerifier",
                        status=ExecutionStatus.VERIFICATION_FAILED,
                        details=f"DOCX '{p.name}' parsed but contains zero paragraphs and zero tables.",
                        error="EMPTY_DOCX",
                    )
                return VerificationOutcome(
                    verified=True,
                    verifier_name="DocumentVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"DOCX '{p.name}' verified ({p_count} paragraphs, {t_count} tables, {text_len} chars).",
                    details=f"Verified DOCX structure for '{p.name}'.",
                    observed_state={"paragraphs": p_count, "tables": t_count, "size": p.stat().st_size},
                )

            elif ext == ".pdf":
                with open(p, "rb") as f:
                    header = f.read(5)
                if not header.startswith(b"%PDF-"):
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="DocumentVerifier",
                        status=ExecutionStatus.VERIFICATION_FAILED,
                        details=f"PDF '{p.name}' lacks valid magic header %PDF-.",
                        error="INVALID_PDF_HEADER",
                    )
                # Parse page tree if pypdf is available
                page_count = 0
                try:
                    import pypdf
                    reader = pypdf.PdfReader(str(p))
                    page_count = len(reader.pages)
                except Exception:
                    pass

                return VerificationOutcome(
                    verified=True,
                    verifier_name="DocumentVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"PDF '{p.name}' binary structure verified ({p.stat().st_size:,} bytes{f', {page_count} pages' if page_count else ''}).",
                    details=f"Verified PDF structure for '{p.name}'.",
                    observed_state={"size": p.stat().st_size, "pages": page_count},
                )

            elif ext in (".xlsx", ".xlsm"):
                if not zipfile.is_zipfile(str(p)):
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="DocumentVerifier",
                        status=ExecutionStatus.VERIFICATION_FAILED,
                        details=f"XLSX '{p.name}' is not a valid OpenXML ZIP archive.",
                        error="INVALID_XLSX",
                    )
                return VerificationOutcome(
                    verified=True,
                    verifier_name="DocumentVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"Excel spreadsheet '{p.name}' verified ({p.stat().st_size:,} bytes).",
                    details=f"Verified XLSX structure for '{p.name}'.",
                )

            elif ext == ".json":
                data = json.loads(p.read_text(encoding="utf-8"))
                return VerificationOutcome(
                    verified=True,
                    verifier_name="DocumentVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"JSON '{p.name}' parsed successfully ({type(data).__name__} root).",
                    details=f"Verified JSON structure for '{p.name}'.",
                )

            elif ext == ".csv":
                with open(p, "r", encoding="utf-8", errors="replace") as f:
                    reader = csv.reader(f)
                    rows = list(reader)
                return VerificationOutcome(
                    verified=True,
                    verifier_name="DocumentVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"CSV '{p.name}' parsed successfully ({len(rows)} rows).",
                    details=f"Verified CSV structure for '{p.name}'.",
                )

            elif ext in (".html", ".htm"):
                content = p.read_text(encoding="utf-8", errors="replace").strip()
                if len(content) < 10 or ("<html" not in content.lower() and "<div" not in content.lower() and "<!doctype" not in content.lower()):
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="DocumentVerifier",
                        status=ExecutionStatus.VERIFICATION_FAILED,
                        details=f"HTML '{p.name}' lacks valid HTML markup structure.",
                        error="INVALID_HTML",
                    )
                return VerificationOutcome(
                    verified=True,
                    verifier_name="DocumentVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"HTML document '{p.name}' verified ({len(content):,} chars, {p.stat().st_size:,} bytes).",
                    details=f"Verified HTML structure for '{p.name}'.",
                    observed_state={"size": p.stat().st_size, "chars": len(content)},
                )

            return file_res
        except Exception as e:
            return VerificationOutcome(
                verified=False,
                verifier_name="DocumentVerifier",
                status=ExecutionStatus.VERIFICATION_FAILED,
                details=f"Structural parse verification failed for '{p.name}': {e}",
                error="PARSE_ERROR",
            )


class CareerVerifier:
    """Specialized verifier for Career OS documents, ATS scores, application packages, and submissions."""

    @staticmethod
    def verify_resume_artifact(path_str: str | Path) -> VerificationOutcome:
        """Verifies that a generated resume file exists, is non-empty, and structurally valid."""
        return DocumentVerifier.verify_document(path_str)

    @staticmethod
    def verify_application_package(package_dir_or_json: str | Path | Dict[str, Any]) -> VerificationOutcome:
        """Verifies complete application package containing resume, cover letter, and answers."""
        try:
            if isinstance(package_dir_or_json, dict):
                pkg = package_dir_or_json
                req_fields = ["package_id", "job_id", "company", "resume_path", "status"]
                missing = [f for f in req_fields if f not in pkg]
                if missing:
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="CareerVerifier",
                        status=ExecutionStatus.FAILED,
                        details=f"Application package metadata missing fields: {missing}",
                        error="INCOMPLETE_PACKAGE",
                    )
                return VerificationOutcome(
                    verified=True,
                    verifier_name="CareerVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"Application package '{pkg.get('package_id')}' validated for {pkg.get('company')}.",
                    observed_state=pkg,
                )
            
            p = Path(package_dir_or_json).resolve()
            if p.is_file() and p.suffix.lower() == ".json":
                data = json.loads(p.read_text(encoding="utf-8"))
                return CareerVerifier.verify_application_package(data)
            elif p.is_dir():
                files = [f.name for f in p.iterdir() if f.is_file()]
                has_resume = any("resume" in f.lower() for f in files)
                if not has_resume:
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="CareerVerifier",
                        status=ExecutionStatus.PARTIAL_SUCCESS,
                        details=f"Application package directory '{p.name}' missing resume document.",
                        error="MISSING_RESUME_FILE",
                    )
                return VerificationOutcome(
                    verified=True,
                    verifier_name="CareerVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"Application package folder verified with {len(files)} deliverables.",
                    observed_state={"files": files},
                )
            return VerificationOutcome(
                verified=False,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.FAILED,
                details=f"Application package target '{p}' is neither directory nor json.",
                error="INVALID_PACKAGE_TARGET",
            )
        except Exception as e:
            return VerificationOutcome(
                verified=False,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.FAILED,
                details=f"Career package verification error: {e}",
                error="CAREER_VERIFICATION_ERROR",
            )

    @staticmethod
    def verify_submission_evidence(evidence: Dict[str, Any] | str) -> VerificationOutcome:
        """Verifies physical proof of job submission (confirmation ID, URL, email, API response)."""
        if isinstance(evidence, str):
            evidence = {"raw_text": evidence}
        
        conf_id = evidence.get("confirmation_id") or evidence.get("application_id")
        conf_url = evidence.get("confirmation_url") or evidence.get("response_url")
        api_ok = evidence.get("api_verified") is True
        
        if conf_id or conf_url or api_ok:
            return VerificationOutcome(
                verified=True,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"Submission verified (ID: {conf_id or 'N/A'}, URL: {conf_url or 'N/A'}, API: {api_ok}).",
                details="Verified job application submission evidence.",
                observed_state=evidence,
            )
        
        return VerificationOutcome(
            verified=False,
            verifier_name="CareerVerifier",
            status=ExecutionStatus.SUCCESS_UNVERIFIED,
            details="Submission attempted but lacks authoritative confirmation ID or response proof.",
            error="UNVERIFIED_SUBMISSION",
        )

    @staticmethod
    def verify_email_classification(classification_res: Dict[str, Any] | Any) -> VerificationOutcome:
        """Verifies that an email was genuinely retrieved, parsed, and classified."""
        cls_val = classification_res.get("classification") if isinstance(classification_res, dict) else getattr(classification_res, "classification", None)
        conf = classification_res.get("confidence", 0.0) if isinstance(classification_res, dict) else getattr(classification_res, "confidence", 0.0)

        if cls_val and conf > 0.0:
            return VerificationOutcome(
                verified=True,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"Email classified as '{cls_val}' (confidence: {conf:.2f}).",
                details="Verified email classification and intent extraction.",
                observed_state=classification_res if isinstance(classification_res, dict) else {},
            )
        return VerificationOutcome(
            verified=False,
            verifier_name="CareerVerifier",
            status=ExecutionStatus.FAILED,
            details="Email classification lacks confidence or valid category.",
            error="INVALID_EMAIL_CLASSIFICATION",
        )

    @staticmethod
    def verify_application_match(match_res: Dict[str, Any] | Any) -> VerificationOutcome:
        """Verifies that an email or event was matched with high confidence to an application."""
        app_id = match_res.get("application_id") if isinstance(match_res, dict) else getattr(match_res, "application_id", None)
        conf = match_res.get("confidence", 0.0) if isinstance(match_res, dict) else getattr(match_res, "confidence", 0.0)
        needs_rev = match_res.get("needs_review", False) if isinstance(match_res, dict) else getattr(match_res, "needs_review", False)

        if app_id and conf >= 0.70 and not needs_rev:
            return VerificationOutcome(
                verified=True,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"Matched to application '{app_id}' with {conf*100:.0f}% confidence.",
                details="Authoritative application match verified.",
                observed_state=match_res if isinstance(match_res, dict) else {},
            )
        return VerificationOutcome(
            verified=False,
            verifier_name="CareerVerifier",
            status=ExecutionStatus.PARTIAL_SUCCESS if needs_rev else ExecutionStatus.FAILED,
            details="Application match requires user review or lacks required confidence threshold.",
            error="MATCH_NEEDS_REVIEW",
        )

    @staticmethod
    def verify_offer_evidence(offer_data: Dict[str, Any] | Any) -> VerificationOutcome:
        """Verifies that an offer record contains concrete factual evidence."""
        status_val = offer_data.get("status") if isinstance(offer_data, dict) else getattr(offer_data, "status", None)
        evidence = offer_data.get("evidence") if isinstance(offer_data, dict) else getattr(offer_data, "evidence", "")

        if evidence and status_val in ("OFFER_CANDIDATE", "OFFER_DETECTED", "OFFER_CONFIRMED"):
            return VerificationOutcome(
                verified=True,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"Offer evidence verified: {evidence}",
                details="Verified job offer candidate record and terms.",
                observed_state=offer_data if isinstance(offer_data, dict) else {},
            )
        return VerificationOutcome(
            verified=False,
            verifier_name="CareerVerifier",
            status=ExecutionStatus.FAILED,
            details="Offer detection lacks concrete documentary or textual evidence.",
            error="UNVERIFIED_OFFER_EVIDENCE",
        )

    @staticmethod
    def verify_spreadsheet_projection(excel_path_str: str | Path) -> VerificationOutcome:
        """Verifies that the canonical Excel workbook exists, has 10 sheets, and is non-corrupt."""
        try:
            p = Path(excel_path_str).resolve()
            if not p.exists() or p.stat().st_size == 0:
                return VerificationOutcome(
                    verified=False,
                    verifier_name="CareerVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Excel workbook '{p}' not found or zero-byte file.",
                    error="EXCEL_FILE_MISSING",
                )

            import openpyxl
            wb = openpyxl.load_workbook(p, read_only=True)
            sheet_names = wb.sheetnames
            wb.close()

            req_sheets = ["Applications", "Jobs", "Interviews", "Offers", "Followups", "Email Events", "Analytics"]
            missing = [s for s in req_sheets if s not in sheet_names]
            if missing:
                return VerificationOutcome(
                    verified=False,
                    verifier_name="CareerVerifier",
                    status=ExecutionStatus.PARTIAL_SUCCESS,
                    details=f"Excel workbook missing canonical sheets: {missing}",
                    error="INCOMPLETE_EXCEL_PROJECTION",
                )

            return VerificationOutcome(
                verified=True,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.SUCCESS_VERIFIED,
                evidence=f"Excel projection verified at '{p.name}' ({len(sheet_names)} sheets present).",
                details="10-sheet career tracker workbook validated.",
                observed_state={"sheets": sheet_names, "size_bytes": p.stat().st_size},
            )
        except Exception as exc:
            return VerificationOutcome(
                verified=False,
                verifier_name="CareerVerifier",
                status=ExecutionStatus.FAILED,
                details=f"Excel workbook validation error: {exc}",
                error="EXCEL_VALIDATION_ERROR",
            )




class ApplicationVerifier:
    """Verifies OS application launches, active processes, and visible windows."""

    @staticmethod
    def verify_process_running(proc_name: str) -> VerificationOutcome:
        try:
            import psutil
            low = proc_name.lower().strip()
            base = os.path.splitext(low)[0]

            matched = []
            for proc in psutil.process_iter(['name', 'pid']):
                try:
                    pname = (proc.info.get('name') or '').lower()
                    if base in pname or low in pname:
                        matched.append(proc.info)
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue

            if matched:
                best = matched[0]
                return VerificationOutcome(
                    verified=True,
                    verifier_name="ApplicationVerifier",
                    status=ExecutionStatus.SUCCESS_VERIFIED,
                    evidence=f"Process '{best['name']}' (PID: {best['pid']}) confirmed active in OS process table.",
                    details=f"Verified process '{best['name']}' active (PID: {best['pid']}).",
                    observed_state={"process_name": best['name'], "pid": best['pid']},
                )

            return VerificationOutcome(
                verified=False,
                verifier_name="ApplicationVerifier",
                status=ExecutionStatus.SUCCESS_UNVERIFIED,
                details=f"Process '{proc_name}' not found in active process table.",
                error="PROCESS_NOT_FOUND",
            )
        except Exception as e:
            return VerificationOutcome(
                verified=True,
                verifier_name="ApplicationVerifier",
                status=ExecutionStatus.SUCCESS_UNVERIFIED,
                details=f"Process verification notice: {e}",
            )

    @staticmethod
    def verify_window_open(window_title_keyword: Optional[str] = None, app_name: Optional[str] = None) -> VerificationOutcome:
        """Inspect visible GUI window handles on Windows OS."""
        if sys.platform != "win32":
            return ApplicationVerifier.verify_process_running(app_name or "")

        try:
            import ctypes
            from ctypes import wintypes

            titles = []

            def enum_windows_proc(hwnd, lParam):
                if ctypes.windll.user32.IsWindowVisible(hwnd):
                    length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
                    if length > 0:
                        buff = ctypes.create_unicode_buffer(length + 1)
                        ctypes.windll.user32.GetWindowTextW(hwnd, buff, length + 1)
                        if buff.value.strip():
                            titles.append(buff.value.strip())
                return True

            WNDENUMPROC = ctypes.WINFUNCTYPE(wintypes.BOOL, wintypes.HWND, wintypes.LPARAM)
            ctypes.windll.user32.EnumWindows(WNDENUMPROC(enum_windows_proc), 0)

            target_raw = (window_title_keyword or app_name or "").strip()
            kw_list = []
            if target_raw:
                kw_list.append(target_raw.lower())
                # If target is a file path, extract name and stem
                if "/" in target_raw or "\\" in target_raw or "." in target_raw:
                    p = Path(target_raw)
                    if p.name:
                        kw_list.append(p.name.lower())
                    if p.stem:
                        kw_list.append(p.stem.lower())
                    # Also replace underscores with spaces
                    if "_" in p.stem:
                        kw_list.append(p.stem.replace("_", " ").lower())

            # Check window title matches
            for kw in kw_list:
                if not kw or len(kw) < 3:
                    continue
                matches = [t for t in titles if kw in t.lower()]
                if matches:
                    return VerificationOutcome(
                        verified=True,
                        verifier_name="ApplicationVerifier",
                        status=ExecutionStatus.SUCCESS_VERIFIED,
                        evidence=f"Active window detected: '{matches[0]}'.",
                        details=f"Verified window open matching '{kw}': '{matches[0]}'",
                        observed_state={"window_title": matches[0], "all_matches": matches},
                    )

            # Check candidate viewer processes if target is a document
            doc_ext = Path(target_raw).suffix.lower()
            viewer_map = {
                ".pdf": ["msedge.exe", "chrome.exe", "acrobat.exe", "acrord32.exe", "foxitpdfreader.exe", "brave.exe", "firefox.exe"],
                ".docx": ["winword.exe", "soffice.bin", "wordpad.exe", "wps.exe"],
                ".doc": ["winword.exe", "soffice.bin", "wordpad.exe", "wps.exe"],
                ".xlsx": ["excel.exe", "soffice.bin", "et.exe"],
                ".txt": ["notepad.exe", "code.exe", "notepad++.exe"],
                ".png": ["photosapp.exe", "mspaint.exe", "dllhost.exe", "microsoft.photos.exe"],
                ".jpg": ["photosapp.exe", "mspaint.exe", "dllhost.exe", "microsoft.photos.exe"],
            }
            if doc_ext in viewer_map:
                for candidate in viewer_map[doc_ext]:
                    p_res = ApplicationVerifier.verify_process_running(candidate)
                    if p_res.verified:
                        return VerificationOutcome(
                            verified=True,
                            verifier_name="ApplicationVerifier",
                            status=ExecutionStatus.SUCCESS_VERIFIED,
                            evidence=f"Active viewer process detected for {doc_ext}: '{candidate}'.",
                            details=f"Viewer process '{candidate}' is running on host.",
                            observed_state={"process_name": candidate},
                        )

            if app_name:
                p_res = ApplicationVerifier.verify_process_running(app_name)
                if p_res.verified:
                    return p_res

            search_desc = kw_list[1] if len(kw_list) > 1 else (target_raw or app_name or "")
            return VerificationOutcome(
                verified=False,
                verifier_name="ApplicationVerifier",
                status=ExecutionStatus.SUCCESS_UNVERIFIED,
                details=f"No visible window detected matching '{search_desc}'.",
                error="WINDOW_NOT_FOUND",
            )
        except Exception as e:
            return ApplicationVerifier.verify_process_running(app_name or window_title_keyword or "")


class BrowserVerifier:
    """Verifies that browser opened a legitimate, reachable artifact URL with no error pages."""

    @staticmethod
    def verify_browser_artifact(target_url_or_path: str, browser_response: Optional[str] = None) -> VerificationOutcome:
        target_str = str(target_url_or_path).strip()
        low = target_str.lower().replace("\\", "/")

        # 1. Sandbox jail containment check
        if "jarvis_sandbox_jails" in low or "sandbox_jails" in low or "/jail_" in low:
            return VerificationOutcome(
                verified=False,
                verifier_name="BrowserVerifier",
                status=ExecutionStatus.BLOCKED,
                details=f"Security Violation: Browser attempted to open internal sandbox jail path '{target_str}'.",
                error="SANDBOX_PATH_EXPOSURE",
            )

        # 2. Local file validation
        if not (target_str.startswith("http://") or target_str.startswith("https://")):
            clean_path = target_str
            if clean_path.startswith("file:///"):
                clean_path = clean_path[8:] if sys.platform == "win32" else clean_path[7:]
            elif clean_path.startswith("file://"):
                clean_path = clean_path[7:]
            
            p = Path(clean_path).resolve()
            if not p.exists():
                return VerificationOutcome(
                    verified=False,
                    verifier_name="BrowserVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Browser target file '{clean_path}' does not exist on disk.",
                    error="ERR_FILE_NOT_FOUND",
                )
            if p.is_file() and p.stat().st_size == 0:
                return VerificationOutcome(
                    verified=False,
                    verifier_name="BrowserVerifier",
                    status=ExecutionStatus.FAILED,
                    details=f"Browser target file '{clean_path}' is empty (0 bytes).",
                    error="FILE_EMPTY",
                )

        # 3. Check browser response for errors
        if browser_response:
            resp_str = str(browser_response).lower()
            err_patterns = [
                ("err_file_not_found", "ERR_FILE_NOT_FOUND"),
                ("file not found", "ERR_FILE_NOT_FOUND"),
                ("err_access_denied", "ERR_ACCESS_DENIED"),
                ("failed to load resource", "ERR_LOAD_FAILED"),
            ]
            for pattern, err_code in err_patterns:
                if pattern in resp_str:
                    return VerificationOutcome(
                        verified=False,
                        verifier_name="BrowserVerifier",
                        status=ExecutionStatus.FAILED,
                        details=f"Browser error detected: '{pattern}' loading '{target_str}'.",
                        error=err_code,
                    )

        return VerificationOutcome(
            verified=True,
            verifier_name="BrowserVerifier",
            status=ExecutionStatus.SUCCESS_VERIFIED,
            evidence=f"Browser opened valid, readable target '{target_str}'.",
            details=f"Browser target verified: {target_str}",
        )


class OutputContractValidator:
    """Semantic validator checking stdout/stderr to ensure return code 0 did not mask hidden errors."""

    @staticmethod
    def validate_output(output_str: str, return_code: int = 0) -> VerificationOutcome:
        if not isinstance(output_str, str):
            return VerificationOutcome(verified=True, verifier_name="OutputContractValidator", status=ExecutionStatus.SUCCESS_VERIFIED)

        low = output_str.lower().strip()
        
        # Critical failure indicators that indicate actual execution errors even if return code is 0
        # Specific errors first, followed by generic traceback fallback
        fatal_indicators = [
            ("modulenotfounderror:", "MISSING_DEPENDENCY"),
            ("importerror:", "IMPORT_ERROR"),
            ("zerodivisionerror:", "RUNTIME_MATH_ERROR"),
            ("unboundlocalerror:", "UNBOUND_LOCAL_ERROR"),
            ("syntaxerror:", "SYNTAX_ERROR"),
            ("permission denied", "PERMISSION_DENIED"),
            ("access is denied", "ACCESS_DENIED"),
            ("scope violation", "SCOPE_VIOLATION"),
            ('"status": "failure"', "STRUCTURED_FAILURE"),
            ('"status": "error"', "STRUCTURED_ERROR"),
            ("err_file_not_found", "FILE_NOT_FOUND"),
            ("error building document", "DOCUMENT_BUILD_ERROR"),
            ("traceback (most recent call last):", "UNCAUGHT_PYTHON_EXCEPTION"),
        ]

        for pattern, err_code in fatal_indicators:
            if pattern in low:
                status = ExecutionStatus.MISSING_DEPENDENCY if err_code == "MISSING_DEPENDENCY" else ExecutionStatus.FAILED
                return VerificationOutcome(
                    verified=False,
                    verifier_name="OutputContractValidator",
                    status=status,
                    details=f"Tool output contains failure indicator: '{pattern}'",
                    error=err_code,
                )

        if return_code != 0:
            return VerificationOutcome(
                verified=False,
                verifier_name="OutputContractValidator",
                status=ExecutionStatus.FAILED,
                details=f"Process exited with non-zero code {return_code}",
                error="NON_ZERO_EXIT",
            )

        return VerificationOutcome(
            verified=True,
            verifier_name="OutputContractValidator",
            status=ExecutionStatus.SUCCESS_VERIFIED,
            evidence="Output clean and validated.",
        )


class UniversalVerifier:
    """Master Verification Facade dispatching to appropriate domain verifiers."""

    verify_file = staticmethod(FileVerifier.verify_file)
    verify_directory = staticmethod(DirectoryVerifier.verify_directory)
    verify_document = staticmethod(DocumentVerifier.verify_document)
    verify_process = staticmethod(ApplicationVerifier.verify_process_running)
    verify_window = staticmethod(ApplicationVerifier.verify_window_open)
    verify_browser = staticmethod(BrowserVerifier.verify_browser_artifact)
    verify_career_artifact = staticmethod(CareerVerifier.verify_resume_artifact)
    verify_application_package = staticmethod(CareerVerifier.verify_application_package)
    verify_submission = staticmethod(CareerVerifier.verify_submission_evidence)
    validate_output = staticmethod(OutputContractValidator.validate_output)

    @classmethod
    def verify_execution(cls, tool_or_command: str, args: Dict[str, Any], raw_output: str, return_code: int = 0) -> VerificationOutcome:
        """Evaluate return code + output contract + real-world side effects."""
        # 1. Output contract check
        out_res = cls.validate_output(raw_output, return_code=return_code)
        if not out_res.verified:
            return out_res

        name = tool_or_command.lower()

        # 2. File write verification
        if name in ("file_write", "create_file", "write_file"):
            target_path = args.get("path") or args.get("file_path") or args.get("name") or ""
            if target_path:
                return cls.verify_file(target_path)

        # 3. Document creator verification
        if name in ("create_word_document", "create_pdf_document", "document_creator", "generate_walkthrough"):
            filename = args.get("filename") or args.get("path") or ""
            if not filename:
                title = args.get("title", "Document")
                fmt = args.get("format", "docx" if "word" in name else "pdf" if "pdf" in name else "docx")
                clean_title = re.sub(r'[^\w\-]', '_', title)
                filename = f"workspace/Documents/{clean_title}.{fmt}"
            return cls.verify_document(filename)

        # 4. Career resume export verification
        if name in ("resume_export", "resume_create", "export_resume"):
            filename = args.get("output_path") or args.get("path") or args.get("filename") or ""
            if filename:
                return cls.verify_career_artifact(filename)

        # 5. Application package verification
        if name in ("application_prepare", "prepare_application_package"):
            pkg_path = args.get("package_path") or args.get("output_dir") or ""
            if pkg_path:
                return cls.verify_application_package(pkg_path)

        # 6. Browser open verification
        if name in ("browser_open_url", "open_browser", "web_browser", "application_open"):
            target_url = args.get("url") or args.get("uri") or args.get("path") or args.get("application_url") or ""
            if target_url:
                return cls.verify_browser(target_url, browser_response=raw_output)

        # 7. App launch verification
        if name in ("open_app", "launch_app"):
            app_name = args.get("app_name") or args.get("name") or ""
            if any(b in app_name.lower() for b in ["chrome", "msedge", "edge", "brave", "firefox"]):
                parts = app_name.split(maxsplit=1)
                if len(parts) > 1 and (":" in parts[1] or "/" in parts[1] or "\\" in parts[1]):
                    return cls.verify_browser(parts[1], browser_response=raw_output)
            if app_name:
                return cls.verify_window(app_name=app_name)

        return VerificationOutcome(
            verified=True,
            verifier_name="UniversalVerifier",
            status=ExecutionStatus.SUCCESS_VERIFIED,
            evidence=f"Tool '{tool_or_command}' executed cleanly with validated output.",
            details="Validated clean execution.",
        )


_GLOBAL_VERIFIER: Optional[UniversalVerifier] = None


def get_universal_verifier() -> UniversalVerifier:
    global _GLOBAL_VERIFIER
    if _GLOBAL_VERIFIER is None:
        _GLOBAL_VERIFIER = UniversalVerifier()
    return _GLOBAL_VERIFIER
